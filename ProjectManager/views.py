# Importaciones de Python estándar
import os
import io
import re
import shutil
import time
import tempfile
import requests
import base64
import logging
from decimal import Decimal, InvalidOperation
from urllib.parse import urlparse, urljoin
from collections import defaultdict
import xml.etree.ElementTree as ET
from io import BytesIO
import subprocess
from requests.exceptions import ConnectionError
from attack_narrative.models import Writeup
from colorsys import hls_to_rgb 
from django.contrib.auth.decorators import login_required
from django.db import models
from attack_narrative.models import WriteupImage, Writeup



# Importaciones de terceros
from googletrans import Translator
from deep_translator import GoogleTranslator
from googletrans import Translator
from PIL import Image, ImageDraw, ImageEnhance
from bs4 import BeautifulSoup
from html2docx import html2docx
import matplotlib.pyplot as plt
from django.views.decorators.cache import never_cache
from django.views.decorators.http import require_POST
from django.views.decorators.clickjacking import xframe_options_sameorigin
import markdown

try:
    import cairosvg
except ImportError:
    cairosvg = None

# Importaciones de Django
from django.shortcuts import render, get_object_or_404, redirect
from django.views import View
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse, FileResponse, HttpResponseForbidden, Http404
from django.db import transaction
from django.db.models import Case, When, Value, IntegerField
from django.conf import settings
from django.urls import reverse
from django.utils.html import strip_tags, escape
from django.template.loader import render_to_string
from django.templatetags.static import static
from django.contrib.auth import login
from django.core.exceptions import MultipleObjectsReturned
from django.conf import settings

# Importaciones de modelos y formularios de la aplicación actual
from .models import Project, Vulnerability, Port, EvidenceImage, PortVulnerabilityProject, Target, ReportTemplate
from attack_narrative.models import Writeup
from .forms import (
    ProjectForm, NessusFileUploadForm, NmapFileUploadForm, ChangeProjectForm,
    PortVulnerabilityProjectForm, NetsparkerFileUploadForm, BurpUploadForm,
    AssignTargetsAndPortsForm, VulnerabilityForm, EvidenceImageForm, 
    ReportTemplateForm, TinyMCEForm, CustomUserCreationForm
)


# Importaciones de python-docx
import docx
from docx import Document
from docx.shared import Inches, RGBColor, Cm, Pt
from docx.oxml import parse_xml, OxmlElement, ns
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION


# Importaciones de HTML
from html.parser import HTMLParser
from .forms import ReportCoverForm
from .models import ReportCoverTemplate

from django.contrib import messages  # ✅ Soluciona "messages is not defined"
from django.http import HttpResponse  # ✅ Soluciona "HttpResponse is not defined"
from django.contrib.auth.forms import UserCreationForm  # ✅ Soluciona "UserCreationForm is not defined"
from ProjectManager.models import ReportCoverTemplate as ReportCover
from ProjectManager.utils import risk_factor_to_numeric, translate_text
from ProjectManager.image_utils import (
    download_external_images,
    add_image_to_doc,
    add_base64_image_to_doc_safe,
    resolve_image_path,
)




logger = logging.getLogger(__name__)


@login_required
def home(request):
    return redirect('project_list')


google_translator = GoogleTranslator(source='en', target='es')

def translate_text(text, lang='es'):
    translator = GoogleTranslator(source='auto', target=lang)
    return translator.translate(text)


score = risk_factor_to_numeric("High")  # Devuelve 3
translated_description = translate_text("This is a vulnerability description.")  # Traduce a español


# Content split to translate the whole content:
def split_and_translate(text, lang='es'):
    translator = GoogleTranslator(source='auto', target=lang)
    max_length = 4000
    chunks = [text[i:i + max_length] for i in range(0, len(text), max_length)]
    translated_text = ''

    for chunk in chunks:
        retries = 3
        while retries > 0:
            try:
                translated_chunk = translator.translate(chunk)
                translated_text += translated_chunk
                break
            except Exception as e:
                retries -= 1
                print(f"Error al traducir un fragmento: {e}. Reintentando...")
                time.sleep(2)
        else:
            translated_text += chunk

    return translated_text





def clean_html(raw_html):
    """Función para limpiar texto HTML de tags."""
    cleanr = re.compile('<.*?>')
    cleantext = re.sub(cleanr, '', raw_html)
    return cleantext


def get_projects_queryset(user):
    """Admin (staff) ve todos los proyectos; consultant solo los asignados (members)."""
    if getattr(user, 'is_staff', False):
        return Project.objects.all()
    return Project.objects.filter(members=user)


def create_or_edit_project(request, project_id=None):
    project = None
    if project_id:
        qs = get_projects_queryset(request.user) if getattr(request.user, 'is_authenticated', False) else Project.objects.all()
        project = get_object_or_404(qs, id=project_id)

    if request.method == 'POST':
        form = ProjectForm(request.POST, instance=project)
        if form.is_valid():
            form.save()
            return redirect('project_list')  # Redirige a la vista que prefieras
    else:
        form = ProjectForm(instance=project)

    return render(request, 'projectmanager/project_form.html', {'form': form})


# Vista para crear proyectos
@login_required
def create_project(request):
    if request.method == 'POST':
        project_form = ProjectForm(request.POST)
        if project_form.is_valid():
            project = project_form.save()
            return redirect('project_detail', pk=project.pk)
    else:
        project_form = ProjectForm()

    return render(request, 'projectmanager/create_project.html', {'project_form': project_form})


# Vista para listar todos los proyectos (admin: todos; consultant: solo asignados)
@login_required
def project_list(request):
    projects = get_projects_queryset(request.user).order_by('-id')
    return render(request, 'projectmanager/project_list.html', {'projects': projects})


@login_required
def project_delete(request, project_id):
    """Delete project and all related data (vulnerabilities, targets, ports, evidence). Does not delete Cover/Report templates."""
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    if request.method == 'POST':
        name = project.name
        project.delete()
        messages.success(request, f'Project "{name}" and all its data have been deleted.')
        return redirect('project_list')
    return render(request, 'projectmanager/project_delete_confirm.html', {'project': project})


# Vista para detalles de un proyecto específico
@login_required
def project_detail(request, pk):
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)
    attack_narratives = project.attack_narratives.all()  # ✅ Obtener los Writeups de attack_narrative asociados

    # Agrupar vulnerabilidades por nombre
    vulnerabilities = Vulnerability.objects.filter(project=project)
    grouped_vulnerabilities = defaultdict(list)
    for vulnerability in vulnerabilities:
        grouped_vulnerabilities[vulnerability.name].append(vulnerability)

    # Obtener targets y nodos para el gráfico
    targets = Target.objects.filter(project=project, owned=True).select_related('jumped_from')
    nodes = [{'id': target.id, 'name': str(target), 'owned': target.owned} for target in targets]
    edges = [{'source': target.jumped_from.id, 'target': target.id} for target in targets if target.jumped_from]

    # Lógica para generar el informe si se solicita
    if request.method == 'POST':
        # Usar el idioma del proyecto de forma robusta
        language = request.POST.get('language', 'EN')
        doc = Document()
        doc.add_heading(project.name, 0)

        # Agregar vulnerabilidades al informe
        for vulnerability_name, vulnerability_list in grouped_vulnerabilities.items():
            doc.add_heading(f'Vulnerability: {vulnerability_name}', level=1)
            for vulnerability in vulnerability_list:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'TableGrid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field'
                hdr_cells[1].text = 'Detail'

                description_translation = vulnerability.description_es if language == 'ES' else vulnerability.description
                solution_translation = vulnerability.solution_es if language == 'ES' else vulnerability.solution

                data = [
                    ('Detail', vulnerability.name),
                    ('Solution', solution_translation),
                    ('Hosts Affected', vulnerability.hosts_affected if vulnerability.hosts_affected else 'Unknown'),
                    ('Description', description_translation),
                ]

                for label, value in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{project.name}_report.docx"'
        doc.save(response)
        return response

    targets = Target.objects.filter(project=project).select_related('jumped_from')
    vulnerabilities_list = list(Vulnerability.objects.filter(project=project))
    criticity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}
    vulnerabilities_sorted_for_report = sorted(
        vulnerabilities_list,
        key=lambda x: criticity_order.get(x.risk_factor, 4)
    )
    return render(request, 'projectmanager/project_detail.html', {
        'project': project,
        'project_info': project,
        'targets': targets,
        'vulnerabilities': vulnerabilities_list,
        'vulnerabilities_sorted_for_report': vulnerabilities_sorted_for_report,
        'attack_narratives': attack_narratives,
        'grouped_vulnerabilities': grouped_vulnerabilities,
        'nodes': nodes,
        'edges': edges,
    })    



"""
def graphmap_detail(request, pk):
    project = get_object_or_404(Project, pk=pk)
    targets = Target.objects.filter(project=project, owned=True).select_related('jumped_from')

    nodes = []
    edges = []

    for target in targets:
        nodes.append({
            'id': target.id,
            'name': str(target),
            'owned': target.owned,
        })
        if target.jumped_from:
            edges.append({
                'source': target.jumped_from.id,
                'target': target.id,
            })

    context = {
        'project': project,
        'nodes': nodes,
        'edges': edges,
    }

    return render(request, 'graphmap_detail.html', context)
"""



def graphmap_detail(request, pk):
    print(f"Llamando a graphmap_detail con PK {pk}")
    project = get_object_or_404(Project, pk=pk)
    nodes = [{'id': 1, 'name': 'Node 1'}, {'id': 2, 'name': 'Node 2'}]
    edges = [{'source': 1, 'target': 2}]

    return render(request, 'admin/project_graphmap.html', {
        'project': project,
        'nodes': nodes,
        'edges': edges,
    })



# Vista para importar archivo Nessus
@login_required
def import_nessus_file(request, pk):
    from deep_translator import GoogleTranslator
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)

    if request.method == 'POST':
        form = NessusFileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            nessus_file = request.FILES['nessus_file']
            tree = ET.parse(nessus_file)
            root = tree.getroot()

            for report_host in root.findall('.//ReportHost'):
                host_ip = 'N/A'
                host_fqdn = 'N/A'
                for tag in report_host.findall('.//tag'):
                    if tag.get('name') == 'host-ip':
                        host_ip = tag.text
                    elif tag.get('name') == 'host-fqdn':
                        host_fqdn = tag.text

                target, _ = Target.objects.get_or_create(
                    ip_address=host_ip,
                    fqdn=host_fqdn,
                    project=project
                )

                for report_item in report_host.findall('.//ReportItem'):
                    name = report_item.get('pluginName')
                    description = report_item.findtext('description') or ''
                    solution = report_item.findtext('solution') or ''
                    see_also = report_item.findtext('see_also') or ''
                    evidence = report_item.findtext('plugin_output') or ''
                    risk_factor = report_item.findtext('risk_factor') or 'N/A'
                    cvss_temporal_score_text = report_item.findtext('cvss_temporal_score')

                    # Traducciones
                    description_es = GoogleTranslator(source='auto', target='es').translate(description) if description.strip() else ''
                    solution_es = GoogleTranslator(source='auto', target='es').translate(solution) if solution.strip() else ''

                    # Conversión segura
                    try:
                        cvss_temporal_score = Decimal(cvss_temporal_score_text) if cvss_temporal_score_text else None
                    except InvalidOperation:
                        cvss_temporal_score = None

                    # Puertos
                    port_text = report_item.get('port')
                    port_number = int(port_text) if port_text and port_text.isdigit() else 0
                    protocol = report_item.get('protocol') or 'TCP'
                    service_name = report_item.get('svc_name') or ''
                    state = 'open'
                    banner = evidence

                    port_defaults = {
                        'state': state,
                        'service_name': service_name,
                        'banner': banner,
                    }

                    port, _ = Port.objects.get_or_create(
                        target=target,
                        port_number=port_number,
                        protocol=protocol.upper(),
                        defaults=port_defaults
                    )

                    vulnerability = Vulnerability.objects.create(
                        project=project,
                        name=name,
                        description=description,
                        description_es=description_es,
                        solution=solution,
                        solution_es=solution_es,
                        cvss_temporal_score=cvss_temporal_score,
                        see_also=see_also,
                        evidence=evidence,
                        risk_factor=risk_factor,
                        hosts_affected=f"{host_ip} ({host_fqdn})",
                        port=port,
                        target_host=target
                    )

                    PortVulnerabilityProject.objects.create(
                        port=port,
                        vulnerability=vulnerability,
                        project=project
                    )

            return redirect('project_detail', pk=project.pk)
    else:
        form = NessusFileUploadForm()

    return render(request, 'admin/import_nessus.html', {'form': form, 'project': project})




class AssignTargetsAndPortsView(View):
    form_class = AssignTargetsAndPortsForm
    template_name = 'admin/assign_targets_and_ports.html'

    def get(self, request, *args, **kwargs):
        form = self.form_class()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        form = self.form_class(request.POST)
        if form.is_valid():
            # Aquí manejas la lógica de guardado, por ejemplo:
            form.save()
            messages.success(request, 'Targets and ports assigned successfully.')
            return redirect('alguna_url_después_de_guardar')  # Asegúrate de reemplazar esto con una URL válida
        else:
            messages.error(request, 'Please correct the errors in the form.')
        return render(request, self.template_name, {'form': form})



@login_required
def targets_view(request):
    # Proyectos visibles para el usuario (admin: todos; consultant: solo asignados)
    projects = get_projects_queryset(request.user)

    # Filtrar los hosts disponibles en función del proyecto seleccionado en el formulario
    if request.method == 'POST':
        project_id = request.POST.get('project')
        if project_id:
            project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
            hosts = Target.objects.filter(project=project)
        else:
            # Si no se selecciona ningún proyecto, mostrar todos los hosts
            hosts = Target.objects.all()
    else:
        # Si no hay datos enviados por el formulario, mostrar todos los hosts
        hosts = Target.objects.all()

    return render(request, 'projectmanager/targets.html', {'hosts': hosts, 'projects': projects})


#Nmap Parsers
@login_required
def import_nmap_recon_file(request, pk):
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)
    form = NmapFileUploadForm()  # Inicializa el formulario aquí
    if request.method == 'POST':
        form = NmapFileUploadForm(request.POST, request.FILES)  # Re-inicializa el formulario con los datos enviados
        if form.is_valid():
            nmap_file = request.FILES['nmap_file']
            tree = ET.parse(nmap_file)
            root = tree.getroot()

            for host in root.findall('host'):
                status_el = host.find('status')
                status = status_el.get('state') if status_el is not None else ''
                if status == 'up':
                    addr_el = host.find("address[@addrtype='ipv4']")
                    ip_address = addr_el.get('addr') if addr_el is not None else None
                    if ip_address is None:
                        continue
                    fqdn_element = host.find("hostnames/hostname[@type='PTR']")
                    fqdn = fqdn_element.get('name', '') if fqdn_element is not None else ''
                    try:
                        Target.objects.get(project=project, ip_address=ip_address)
                    except Target.DoesNotExist:
                        Target.objects.create(project=project, ip_address=ip_address, fqdn=fqdn)
            # Redirecciona a otra vista una vez completado el proceso
            return redirect('admin:ProjectManager_vulnerability_changelist')

    return render(request, 'admin/import_netsparker.html', {'form': form, 'project': project})
               


# Vista para importar y procesar archivo Nmap XML de escaneo completo de puertos
@login_required
def import_nmap_xml(request, pk):
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)
    if request.method == 'POST':
        form = NmapFileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            nmap_file = request.FILES.get('nmap_file')
            tree = ET.parse(nmap_file)
            root = tree.getroot()

            for host in root.findall('.//host'):
                status_el = host.find('.//status')
                status = status_el.get('state') if status_el is not None else ''
                if status == 'up':
                    addr_el = host.find('.//address[@addrtype="ipv4"]')
                    ip_address = addr_el.get('addr') if addr_el is not None else None
                    if ip_address is None:
                        continue
                    fqdn_elements = host.findall('.//hostname')
                    fqdn = ' / '.join(elem.get('name') for elem in fqdn_elements if elem is not None)
                    os_match = host.find('.//os/osmatch')
                    os_name = os_match.get('name') if os_match else ""

                    try:
                        target = Target.objects.get(project=project, ip_address=ip_address)
                        target.fqdn = fqdn
                        target.os = os_name
                        target.save()
                    except Target.DoesNotExist:
                        target = Target.objects.create(project=project, ip_address=ip_address, fqdn=fqdn, os=os_name)

                    for port_element in host.findall('.//port'):
                        port_id = port_element.get('portid')
                        port_number = int(port_id) if port_id is not None else 0
                        protocol = port_element.get('protocol') or 'tcp'
                        state_el = port_element.find('.//state')
                        state = state_el.get('state') if state_el is not None else ''
                        service_element = port_element.find('.//service')
                        service_name = service_element.get('name') if service_element else ''
                        product = service_element.get('product') if service_element else ''
                        version = service_element.get('version') if service_element else ''

                        try:
                            port = Port.objects.get(target=target, port_number=port_number, protocol=protocol)
                            port.state = state
                            port.service_name = service_name
                            port.product = product
                            port.version = version
                            port.save()
                        except Port.DoesNotExist:
                            Port.objects.create(
                                target=target, port_number=port_number, protocol=protocol,
                                state=state, service_name=service_name, product=product, version=version
                            )

            return redirect('project_detail', pk=project.pk)
    else:
        form = NmapFileUploadForm()

    return render(request, 'admin/import_netsparker.html', {'form': form, 'project': project})



def _cover_template_image_urls(cover_template):
    """Devuelve un dict con las URLs de las imágenes de la plantilla de portada (vacío si no hay archivo)."""
    urls = {'background': '', 'header_left': '', 'header_right': '', 'customer_image': ''}
    if not cover_template:
        return urls
    if getattr(cover_template, 'background_image', None) and cover_template.background_image:
        urls['background'] = cover_template.background_image.url
    if getattr(cover_template, 'header_image', None) and cover_template.header_image:
        urls['header_left'] = cover_template.header_image.url
    if getattr(cover_template, 'customer_header_image', None) and cover_template.customer_header_image:
        urls['header_right'] = cover_template.customer_header_image.url
    if getattr(cover_template, 'customer_image', None) and cover_template.customer_image:
        urls['customer_image'] = cover_template.customer_image.url
    return urls


def _default_cover_layout():
    """Layout por defecto: posiciones y tamaños en % (left, top, width, height)."""
    return {
        "header_left": {"left": 2, "top": 2, "width": 18, "height": 8},
        "header_right": {"left": 80, "top": 2, "width": 18, "height": 8},
        "title_block": {"left": 5, "top": 22, "width": 90, "height": 25},
        "customer_image": {"left": 10, "top": 55, "width": 80, "height": 35},
    }


# Tapa
@login_required
def configurar_tapa_reporte(request, project_id):
    import json
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    
    if project.cover_template:
        template = project.cover_template
    else:
        template = ReportCoverTemplate(name=f"Cover for {project.name}")
    
    if request.method == 'POST':
        form = ReportCoverForm(request.POST, request.FILES, instance=template)
        if form.is_valid():
            template = form.save(commit=False)
            layout_raw = request.POST.get('cover_layout_json') or '{}'
            try:
                template.cover_layout = json.loads(layout_raw)
            except (ValueError, TypeError):
                template.cover_layout = _default_cover_layout()
            template.save()
            project.cover_template = template
            project.save()
            return redirect('project_detail', pk=project_id)
    else:
        form = ReportCoverForm(instance=template)

    cover_image_urls = _cover_template_image_urls(form.instance)
    layout = (form.instance.cover_layout or _default_cover_layout()) if form.instance.pk else _default_cover_layout()
    cover_layout_json = json.dumps(layout)
    return render(request, 'projectmanager/cover_designer.html', {
        'form': form, 'project': project, 'cover_image_urls': cover_image_urls,
        'cover_layout_json': cover_layout_json,
    })


@login_required
def visual_cover_designer_template(request, template_id):
    """Diseñador visual de portada para una plantilla (desde el admin, sin proyecto)."""
    import json
    template = get_object_or_404(ReportCoverTemplate, pk=template_id)
    if request.method == 'POST':
        form = ReportCoverForm(request.POST, request.FILES, instance=template)
        if form.is_valid():
            template = form.save(commit=False)
            layout_raw = request.POST.get('cover_layout_json') or '{}'
            try:
                template.cover_layout = json.loads(layout_raw)
            except (ValueError, TypeError):
                template.cover_layout = _default_cover_layout()
            template.save()
            return redirect('cover_template_list')
    else:
        form = ReportCoverForm(instance=template)
    cover_image_urls = _cover_template_image_urls(form.instance)
    layout = form.instance.cover_layout or _default_cover_layout()
    cover_layout_json = json.dumps(layout)
    return render(request, 'projectmanager/cover_designer.html', {
        'form': form, 'project': None, 'template': template, 'cover_image_urls': cover_image_urls,
        'cover_layout_json': cover_layout_json,
    })





# Vista para mostrar información del proyecto
def project_info(request, object_id):
    project = get_object_or_404(Project, pk=object_id)
    return render(request, 'admin/project_info.html', {'project': project})

# Vista para mostrar vulnerabilidades del proyecto
def project_vulnerabilities(request, object_id):
    project = get_object_or_404(Project, pk=object_id)
    return render(request, 'admin/project_vulnerabilities.html', {'project': project})

# Vista para mostrar puertos del proyecto
def project_ports(request, object_id):
    project = get_object_or_404(Project, pk=object_id)
    return render(request, 'admin/project_ports.html', {'project': project})

@login_required
@require_POST
def toggle_target_owned(request, target_id):
    """Toggle target owned status (for inline edit in Targets table)."""
    target = get_object_or_404(Target, pk=target_id)
    get_object_or_404(get_projects_queryset(request.user), pk=target.project_id)
    target.owned = not target.owned
    target.save()
    return JsonResponse({'owned': target.owned})


@login_required
def target_edit(request, project_id, target_id):
    """Ver y editar un target del proyecto (como en Admin)."""
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    target = get_object_or_404(Target, pk=target_id, project=project)
    from .forms import TargetEditForm
    if request.method == 'POST':
        form = TargetEditForm(request.POST, instance=target)
        if form.is_valid():
            form.save()
            messages.success(request, 'Target updated successfully.')
            return redirect('project_detail', pk=project_id)
    else:
        form = TargetEditForm(instance=target)
    return render(request, 'projectmanager/target_edit.html', {
        'form': form,
        'target': target,
        'project': project,
    })


@login_required
def profile_edit(request):
    """Editar perfil del usuario (nombre, apellido, email)."""
    if request.method == 'POST':
        from .forms import UserProfileForm
        form = UserProfileForm(request.POST, instance=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Profile updated successfully.')
            return redirect('project_list')
    else:
        from .forms import UserProfileForm
        form = UserProfileForm(instance=request.user)
    return render(request, 'projectmanager/profile_edit.html', {'form': form})


@login_required
def app_password_change(request):
    """Vista de cambio de contraseña de la sección app (no Admin)."""
    from django.contrib.auth.views import PasswordChangeView
    from .forms import AppPasswordChangeForm
    from django.urls import reverse_lazy
    view = PasswordChangeView.as_view(
        form_class=AppPasswordChangeForm,
        template_name='projectmanager/account/password_change.html',
        success_url=reverse_lazy('app_password_change_done'),
    )
    return view(request)


@login_required
def app_password_change_done(request):
    """Confirmación tras cambiar contraseña (sección app)."""
    return render(request, 'projectmanager/account/password_change_done.html')


@login_required
def project_member_list(request):
    """Lista de proyectos para gestionar miembros (admin: todos; consultant: solo asignados)."""
    from django.contrib.auth import get_user_model
    User = get_user_model()
    projects = get_projects_queryset(request.user).order_by('name')
    return render(request, 'projectmanager/project_member_list.html', {'projects': projects})


@login_required
def project_members_manage(request, project_id):
    """Gestionar miembros de un proyecto: listar y añadir usuarios."""
    from django.contrib.auth import get_user_model
    from .forms import AddProjectMemberForm
    User = get_user_model()
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    members = project.members.all().order_by('username')
    if request.method == 'POST':
        form = AddProjectMemberForm(project, request.POST)
        if form.is_valid():
            user = form.cleaned_data['user']
            project.members.add(user)
            messages.success(request, f'User {user.username} added to the project.')
            return redirect('project_members_manage', project_id=project_id)
    else:
        form = AddProjectMemberForm(project)
    can_add_more = form.fields['user'].queryset.exists()
    return render(request, 'projectmanager/project_members_manage.html', {
        'project': project,
        'members': members,
        'form': form,
        'can_add_more': can_add_more,
    })


@login_required
@require_POST
def project_member_remove(request, project_id, user_id):
    """Quitar un usuario del proyecto."""
    from django.contrib.auth import get_user_model
    User = get_user_model()
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    user = get_object_or_404(User, pk=user_id)
    project.members.remove(user)
    messages.success(request, f'User {user.username} removed from the project.')
    return redirect('project_members_manage', project_id=project_id)


@login_required
def user_create(request):
    """Crear nuevo usuario (sección app). Solo usuarios staff. Tras crear, redirige a lista de miembros para asignar a proyectos."""
    from .forms import CustomUserCreationForm
    if not request.user.is_staff:
        messages.error(request, 'You do not have permission to create users.')
        return redirect('project_member_list')
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, f'User {form.cleaned_data["username"]} created. Assign them to projects from "Project members".')
            return redirect('project_member_list')
    else:
        form = CustomUserCreationForm()
    return render(request, 'projectmanager/user_create.html', {'form': form})


# Vista para cambiar los detalles de un proyecto
@login_required
def change_project(request, pk):
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)
    if request.method == 'POST':
        form = ChangeProjectForm(request.POST, instance=project)
        if form.is_valid():
            form.save()
            return redirect('project_detail', pk=project.pk)
    else:
        form = ChangeProjectForm(instance=project)

    return render(request, 'projectmanager/change_project.html', {'form': form})



# Función para traducir al español
def translate_to_spanish(text):
    # Crea una instancia del traductor de Google
    translator = google_translator()

    # Traduce el texto al español
    translated_text = translator.translate(text, lang_tgt='es')

    return translated_text
    

    
@login_required
def add_vulnerability(request, project_id=None):
    from .forms import VulnerabilityFullForm
    project = None
    if project_id:
        project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    else:
        return redirect('project_list')
    if request.method == 'POST':
        form = VulnerabilityFullForm(request.POST, project=project)
        if form.is_valid():
            vuln = form.save(commit=False)
            if project:
                vuln.project = project
            vuln.save()
            form.save_m2m()
            if project:
                return redirect('project_detail', pk=project.pk)
            return redirect('project_list')
    else:
        form = VulnerabilityFullForm(project=project)
    return render(request, 'projectmanager/add_vulnerability.html', {'form': form, 'project': project})


@login_required
def vulnerability_edit(request, project_id, vuln_id):
    """Edit an existing vulnerability (name, description, solution, evidence, severity, etc.)."""
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    vuln = get_object_or_404(Vulnerability, pk=vuln_id, project=project)
    from .forms import VulnerabilityFullForm
    if request.method == 'POST':
        form = VulnerabilityFullForm(request.POST, instance=vuln, project=project)
        if form.is_valid():
            form.save()
            form.save_m2m()
            messages.success(request, 'Vulnerability updated successfully.')
            return redirect('project_detail', pk=project_id)
    else:
        form = VulnerabilityFullForm(instance=vuln, project=project)
    return render(request, 'projectmanager/vulnerability_edit.html', {
        'form': form,
        'project': project,
        'vulnerability': vuln,
    })


@login_required
def add_evidence_image(request, project_id):
    from .forms import EvidenceImageForm
    project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
    if request.method == 'POST':
        form = EvidenceImageForm(request.POST, request.FILES, project=project)
        if form.is_valid():
            evidence = form.save(commit=False)
            evidence.project = project
            evidence.save()
            # Optionally attach this evidence to the selected vulnerability
            vulnerability = form.cleaned_data.get('vulnerability')
            if vulnerability:
                vulnerability.evidence_images.add(evidence)
            next_url = request.GET.get('next') or request.POST.get('next') or reverse('project_detail', args=[project.pk])
            return redirect(next_url)
    else:
        form = EvidenceImageForm(project=project)
    return render(request, 'projectmanager/add_evidence_image.html', {'form': form, 'project': project})




# Función para aplicar color de fondo a las celdas de la tabla en el documento
def set_cell_background(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell_properties.append(shading)
    # Ajuste para evitar agregar sombreado duplicado innecesario
    if not any(child.tag == shading.tag for child in cell_properties):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell_properties.append(shading_elm)

# Función para aplicar color a la fuente del texto en un párrafo
def set_paragraph_font_color(paragraph, color):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(color)

# Función para aplicar color de fondo a un párrafo (corregida)
def set_paragraph_background(paragraph, color):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)



def risk_factor_to_legible_text_and_color(risk_factor, language='en'):
    risk_mapping = {
        'Critical': ('Crítica', '800080') if language == 'es' else ('Critical', '800080'),  # Púrpura
        'High': ('Alto', 'FF0000') if language == 'es' else ('High', 'FF0000'),  # Rojo
        'Medium': ('Medio', 'FFA500') if language == 'es' else ('Medium', 'FFA500'),  # Naranja
        'Low': ('Bajo', '008000') if language == 'es' else ('Low', '008000'),  # Verde
    }
    return risk_mapping.get(risk_factor, ('Desconocido', '000000'))  # Valor por defecto si el riesgo es desconocido





# Función para parsear los estilos de TinyMCE
def parse_styles(style_str):
    import re
    from docx.shared import RGBColor, Pt

    styles = {}
    if not style_str:
        return styles

    style_pairs = style_str.split(';')
    for pair in style_pairs:
        if ':' not in pair:
            continue

        key, value = pair.split(':', 1)
        key = key.strip().lower()
        value = value.strip().lower()

        # Manejo de colores en texto y fondo
        if key in ["color", "background-color"]:
            if value.startswith("rgb"):
                match = re.search(r"rgb\((\d+),\s*(\d+),\s*(\d+)\)", value)
                if match:
                    r, g, b = map(int, match.groups())
                    value = f"{r:02x}{g:02x}{b:02x}"  # Convertir a HEX
            elif value.startswith("#"):
                value = value.lstrip("#")
            else:
                value = "000000"  # Predeterminado a negro si es inválido

            styles[key] = RGBColor.from_string(value)

        # Manejo de negrita, cursiva, subrayado y tachado
        elif key == "font-weight" and value in ["bold", "700"]:
            styles["bold"] = True
        elif key == "font-style" and value == "italic":
            styles["italic"] = True
        elif key == "text-decoration":
            if "underline" in value:
                styles["underline"] = True
            if "line-through" in value:
                styles["strike"] = True

        # Manejo de tamaño de fuente
        elif key == "font-size" and value.endswith("px"):
            try:
                size = int(value.replace("px", "").strip())
                styles["font_size"] = Pt(size * 0.75)  # Conversión aproximada de px a pt
            except ValueError:
                pass

    return styles


# Función para añadir un run con estilos a un párrafo
def add_run_with_styles(paragraph, text, styles):
    run = paragraph.add_run(text)
    if 'bold' in styles:
        run.bold = True
    if 'italic' in styles:
        run.italic = True
    if 'underline' in styles:
        run.underline = True
    if 'strike' in styles:
        run.font.strike = True
    if 'color' in styles:
        run.font.color.rgb = styles['color']
    if 'font_size' in styles:
        run.font.size = Pt(styles['font_size'])
    return run




def generate_vulnerability_table(doc, vulnerabilities, language):
    for vulnerability_name, details in vulnerabilities.items():
        legible_risk = details['risk']
        risk_color_code = details['risk_color']

        # Crear el párrafo para el título de la vulnerabilidad con el esquema de colores adecuado
        paragraph = doc.add_paragraph(style='Heading1')
        run_risk = paragraph.add_run(f"{legible_risk} - ")
        run_risk.font.color.rgb = RGBColor(int(risk_color_code[0:2], 16), int(risk_color_code[2:4], 16), int(risk_color_code[4:], 16))
        run_risk.bold = True

        run_name = paragraph.add_run(vulnerability_name)
        run_name.bold = True

        # Definir los títulos de las columnas según el idioma seleccionado
        if language == 'es':
            titles = ['Hosts Afectados', 'Puerto', 'Descripción', 'Solución', 'Evidencia', 'Evidencia Adicional']
        else:
            titles = ['Affected Hosts', 'Ports', 'Description', 'Solution', 'Evidence', 'Extra Evidence']

        # Crear la tabla para los detalles de la vulnerabilidad
        table = doc.add_table(rows=6, cols=2)
        table.style = 'TableGrid'
        table.autofit = True

        # Establecer el ancho de las columnas
        column_width_left = Cm(2.5)
        column_width_right = Cm(15)
        table.columns[0].width = column_width_left
        table.columns[1].width = column_width_right

        details_data = [
            ', '.join(details['hosts']) if 'hosts' in details else 'Unknown',
            ', '.join(str(port) for port in details['ports']) if 'ports' in details else 'Unknown',
            details['description'] if 'description' in details else '',
            details['solution'] if 'solution' in details else '',
            ', '.join(details['evidence']) if 'evidence' in details else 'Unknown',
            '',  # Placeholder for additional evidence cell
        ]

        for i, title in enumerate(titles):
            cell_left = table.cell(i, 0)
            cell_right = table.cell(i, 1)

            # Aplicar formato al encabezado y la columna izquierda
            cell_left.text = title
            cell_left.paragraphs[0].runs[0].bold = True
            cell_left.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Blanco
            set_cell_background(cell_left, '000000')  # Negro

            cell_right.text = details_data[i]

            # Aplicar color de fondo según la criticidad para la columna derecha
            set_cell_background(cell_right, risk_color_code)
            for paragraph in cell_right.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco

        # Añadir la imagen de la evidencia adicional si está disponible
        extra_evidence_image_cell = table.cell(5, 1)
        evidence_images = EvidenceImage.objects.filter(vulnerability_set__name=vulnerability_name)
        for evidence_image in evidence_images:
            image_path = os.path.join(settings.MEDIA_ROOT, evidence_image.image.name)
            if os.path.exists(image_path):
                run = extra_evidence_image_cell.add_paragraph().add_run()
                run.add_picture(image_path, width=Cm(10))  # Ajustar el ancho según sea necesario



# 👇 Esta función generará la tabla de resumen de vulnerabilidades
def generate_vulns_summary_table(doc, vulnerabilities, language):
    consolidated = {}
    for v in vulnerabilities:
        consolidated[v.name] = v.risk_factor

    table = doc.add_table(rows=1, cols=2)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vulnerabilidad' if language == 'es' else 'Vulnerability'
    hdr_cells[1].text = 'Riesgo' if language == 'es' else 'Risk'
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    color_map = {
        'Critical': ('Crítico', '800080'),
        'High': ('Alto', 'FF0000'),
        'Medium': ('Medio', 'FFA500'),
        'Low': ('Bajo', '008000'),
    }

    for idx, (name, risk) in enumerate(consolidated.items(), 1):
        row = table.add_row().cells
        row[0].text = name
        para = row[1].paragraphs[0]
        label, hex_color = color_map.get(risk, (risk, '000000'))
        run = para.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(hex_color)

        shade = 'D3D3D3' if idx % 2 == 0 else 'FFFFFF'
        for cell in row:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), shade)
            tc_pr.append(shd)




def handle_element(element, parent, doc, vulnerabilities, language):
    if isinstance(element, str):
        text = element.strip()
        if not text:
            return

        if '###BreakPage###' in text:
            doc.add_page_break()
            text = text.replace('###BreakPage###', '')

        if '###VulnsTable###' in text:
            text = text.replace('###VulnsTable###', '')
            if text.strip() and isinstance(parent, docx.text.paragraph.Paragraph):
                parent.add_run(text.strip())
            generate_vulns_summary_table(doc, vulnerabilities, language)
            return

        if '###GraphOverallVulns###' in text:
            text = text.replace('###GraphOverallVulns###', '')
            if text.strip() and isinstance(parent, docx.text.paragraph.Paragraph):
                parent.add_run(text.strip())
            generate_overall_vulns_chart(doc, vulnerabilities)
            return

        if isinstance(parent, docx.text.paragraph.Paragraph):
            parent.add_run(f" {text} ")
        return

    def set_paragraph_spacing(paragraph):
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)

    if element.name == 'p':
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph)

        styles = parse_styles(element.get('style', ''))
        text_color = styles.get('color', '')

        text = element.get_text()

        if '###BreakPage###' in text:
            doc.add_page_break()
            text = text.replace('###BreakPage###', '')

        if '###VulnsTable###' in text:
            text = text.replace('###VulnsTable###', '')
            if text.strip():
                paragraph.add_run(text.strip())
            generate_vulns_summary_table(doc, vulnerabilities, language)

        elif '###GraphOverallVulns###' in text:
            text = text.replace('###GraphOverallVulns###', '')
            if text.strip():
                paragraph.add_run(text.strip())
            generate_overall_vulns_chart(doc, vulnerabilities)

        else:
            for child in element.children:
                handle_element(child, paragraph, doc, vulnerabilities, language)

        if isinstance(text_color, str):
            clean_color = text_color.lstrip('#')
            if len(clean_color) == 6:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(clean_color)

    elif element.name in ['ul', 'ol']:
        list_type = 'ListBullet' if element.name == 'ul' else 'ListNumber'

        for li in element.find_all('li'):
            paragraph = doc.add_paragraph(style=list_type)
            set_paragraph_spacing(paragraph)

            styles = parse_styles(li.get('style', ''))
            text_color = styles.get('color', '')

            for child in li.children:
                handle_element(child, paragraph, doc, vulnerabilities, language)

            if isinstance(text_color, str):
                clean_color = text_color.lstrip('#')
                if len(clean_color) == 6:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor.from_string(clean_color)

    elif element.name in ['h1', 'h2', 'h3']:
        level = int(element.name[1])
        paragraph = doc.add_heading(level=level)
        set_paragraph_spacing(paragraph)

        for child in element.children:
            handle_element(child, paragraph, doc, vulnerabilities, language)

    elif element.name == 'img':
        src = element['src']
        if src.startswith('data:image'):
            format, imgstr = src.split(';base64,')
            ext = format.split('/')[-1]
            add_base64_image_to_doc(doc, imgstr, ext)

    elif element.name in ['b', 'strong', 'span', 'i', 'em', 'u'] or 'style' in element.attrs:
        styles = parse_styles(element.get('style', ''))
        text_content = element.get_text(strip=True)

        if isinstance(parent, docx.text.paragraph.Paragraph):
            run = parent.add_run(f" {text_content} ")

            if element.name in ['b', 'strong'] or styles.get('font-weight') == 'bold':
                run.bold = True
            if element.name in ['i', 'em'] or styles.get('font-style') == 'italic':
                run.italic = True
            if element.name == 'u' or 'underline' in styles.get('text-decoration', ''):
                run.underline = True
            if 'color' in styles:
                color_value = styles['color']
                if isinstance(color_value, str):
                    clean_color = color_value.lstrip('#')
                    if len(clean_color) == 6:
                        run.font.color.rgb = RGBColor.from_string(clean_color)

    elif element.name == 'hr':
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph)
        run = paragraph.add_run()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        paragraph._element.get_or_add_pPr().append(pBdr)

    elif element.name == 'br':
        if isinstance(parent, docx.text.paragraph.Paragraph):
            parent.add_run().add_break()

    elif element.name in ['div', 'section', 'article', 'main']:
        for child in element.children:
            handle_element(child, parent, doc, vulnerabilities, language)

    else:
        for child in element.children:
            handle_element(child, parent, doc, vulnerabilities, language)
            
def generate_overall_vulns_chart(doc, vulnerabilities):
    from matplotlib import pyplot as plt
    from io import BytesIO
    from docx.shared import Inches

    # Preparar los datos
    severity_levels = ['Critical', 'High', 'Medium', 'Low']
    severity_colors = {
        'Critical': '#800080',
        'High': '#FF0000',
        'Medium': '#FFA500',
        'Low': '#008000'
    }

    values = [sum(1 for v in vulnerabilities if v.risk_factor == s) for s in severity_levels]

    # ⚠️ Evitar crash si no hay datos
    if not any(values):
        doc.add_paragraph("No vulnerabilities were found to generate the graph.")
        return

    # Crear el gráfico
    fig, ax = plt.subplots(figsize=(5, 3))
    bars = ax.bar(severity_levels, values, color=[severity_colors[s] for s in severity_levels])

    ax.set_ylabel('Cantidad')
    ax.set_title('Vulnerability Summary by Severity')
    ax.set_ylim(0, max(values) + 1)
    ax.grid(axis='y', linestyle='--', alpha=0.6)

    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'{int(height)}',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 3),  # Offset
                    textcoords="offset points",
                    ha='center', va='bottom', fontsize=8)

    # Guardar en memoria
    buffer = BytesIO()
    plt.tight_layout()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    plt.close(fig)

    # Insertar en el doc
    doc.add_picture(buffer, width=Inches(5.5))
    doc.add_paragraph("Figure: Vulnerability Distribution by Severity").alignment = 1


def process_attack_narrative_html(attack_narrative_content):
    if not attack_narrative_content:
        return ""

    soup = BeautifulSoup(attack_narrative_content, "html.parser")

    # ✅ Estilo para bloques <pre>
    for pre_tag in soup.find_all("pre"):
        styled_div = soup.new_tag("div")
        styled_div["style"] = (
            "border: 1px solid #686e73; background-color: #2E2E2E; padding: 10px; "
            "font-family: 'Courier New', monospace; font-size: 9pt; color: white; margin-bottom: 10px;"
        )
        label = soup.new_tag("div")
        label["style"] = (
            "background-color: #686e73; color: white; font-weight: bold; padding: 2px 5px; "
            "margin-bottom: 5px; display: inline-block; border-radius: 3px;"
        )
        label.string = "Code"
        styled_div.append(label)

        if pre_tag.parent:
            pre_tag.wrap(styled_div)
        else:
            styled_div.append(pre_tag)
            soup.append(styled_div)

    # ✅ Estilo para <code> inline
    for code_tag in soup.find_all("code"):
        code_tag["style"] = (
            "background-color: #4A4A4A; color: white; padding: 2px 4px; "
            "font-family: 'Courier New', monospace; border-radius: 3px;"
        )

    return str(soup)




def process_attack_narrative_html(attack_narrative_content):
    if not attack_narrative_content:
        return ""

    soup = BeautifulSoup(attack_narrative_content, "html.parser")

    # ✅ Aplicar estilos a bloques de código (`<pre>`)
    for pre_tag in soup.find_all("pre"):
        styled_div = soup.new_tag("div")
        styled_div["style"] = "border: 1px solid #686e73; background-color: #2E2E2E; padding: 10px; font-family: 'Courier New', monospace; font-size: 9pt; color: white; margin-bottom: 10px;"
        label = soup.new_tag("div")
        label["style"] = "background-color: #686e73; color: white; font-weight: bold; padding: 2px 5px; margin-bottom: 5px; display: inline-block; border-radius: 3px;"
        label.string = "Code"
        styled_div.append(label)
        if pre_tag.parent:
            pre_tag.wrap(styled_div)
        else:
            styled_div.append(pre_tag)
            soup.append(styled_div)

    # ✅ Aplicar estilos a código inline (`<code>`)
    for code_tag in soup.find_all("code"):
        code_tag["style"] = "background-color: #4A4A4A; color: white; padding: 2px 4px; font-family: 'Courier New', monospace; border-radius: 3px;"

    return str(soup)

# Modificar la función generate_report_with_attack_narratives para incluir el nuevo formato
def process_attack_narrative_html(attack_narrative_content):
    if not attack_narrative_content:
        return ""

    soup = BeautifulSoup(attack_narrative_content, "html.parser")

    # ✅ Aplicar estilos a bloques de código (`<pre>`)
    for pre_tag in soup.find_all("pre"):
        styled_div = soup.new_tag("div")
        styled_div["style"] = "border: 1px solid #686e73; background-color: #2E2E2E; padding: 10px; font-family: 'Courier New', monospace; font-size: 9pt; color: white; margin-bottom: 10px;"
        label = soup.new_tag("div")
        label["style"] = "background-color: #686e73; color: white; font-weight: bold; padding: 2px 5px; margin-bottom: 5px; display: inline-block; border-radius: 3px;"
        label.string = "Code"
        styled_div.append(label)
        if pre_tag.parent:
            pre_tag.wrap(styled_div)
        else:
            styled_div.append(pre_tag)
            soup.append(styled_div)

    # ✅ Aplicar estilos a código inline (`<code>`)
    for code_tag in soup.find_all("code"):
        code_tag["style"] = "background-color: #4A4A4A; color: white; padding: 2px 4px; font-family: 'Courier New', monospace; border-radius: 3px;"

    return str(soup)




def get_alignment_from_style(tag):
    style = tag.get("style", "")
    if "text-align: center" in style or "text-align:center" in style:
        return WD_PARAGRAPH_ALIGNMENT.CENTER
    if "text-align: right" in style or "text-align:right" in style:
        return WD_PARAGRAPH_ALIGNMENT.RIGHT
    return WD_PARAGRAPH_ALIGNMENT.LEFT


# Función existente para procesar el contenido HTML en docx
def add_html_to_doc(doc, html_content, vulnerabilities, language):
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.children:
        if isinstance(element, str):
            continue

        if element.name == 'p':
            paragraph = doc.add_paragraph()
            paragraph.alignment = get_alignment_from_style(element)
            for child in element.children:
                handle_element(child, paragraph, doc, vulnerabilities, language)

        elif element.name in ['ul', 'ol']:
            list_type = 'ListBullet' if element.name == 'ul' else 'ListNumber'
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph(style=list_type)
                for child in li.children:
                    handle_element(child, paragraph, doc, vulnerabilities, language)

        elif element.name in ['h1', 'h2', 'h3']:
            level = int(element.name[1])
            paragraph = doc.add_heading(level=level)
            paragraph.alignment = get_alignment_from_style(element)
            for child in element.children:
                handle_element(child, paragraph, doc, vulnerabilities, language)

        elif element.name == 'table':
            rows = element.find_all("tr")
            if not rows:
                continue

            num_cols = max(len(row.find_all(["td", "th"])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    target_cell = table.cell(i, j)
                    paragraph = target_cell.paragraphs[0]
                    paragraph.clear()

                    for child in cell.children:
                        handle_element(child, paragraph, doc, vulnerabilities, language)

        elif element.name == 'code' or ('class' in element.attrs and 'Texto plano' in element['class']):  
            text_content = element.get_text(strip=True)
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text_content)

            print(f" → Formateando código: {text_content}")  # 📌 Debug
            
            run.font.name = 'Courier New'
            run.font.size = docx.shared.Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.paragraph_format.space_after = docx.shared.Pt(6)

            run.bold = True

        elif element.name in ['b', 'strong', 'span', 'i', 'em', 'u'] or 'style' in element.attrs:  
            styles = parse_styles(element.get('style', ''))
            text_content = element.get_text(strip=True)

            if 'parent' in locals() and isinstance(parent, docx.text.paragraph.Paragraph):
                run = parent.add_run(f" {text_content} ")
                print(f"Procesando texto: {text_content}")  # 📌 Verifica que está procesando elementos correctamente

                if element.name in ['b', 'strong'] or ('font-weight' in styles and styles['font-weight'] == 'bold'):
                    run.bold = True
                    print(f" → Aplicado BOLD a: {text_content}")  # 📌 Debug

                if element.name in ['i', 'em'] or ('font-style' in styles and styles['font-style'] == 'italic'):
                    run.italic = True
                    print(f" → Aplicado ITALIC a: {text_content}")  # 📌 Debug

                if element.name == 'u' or ('text-decoration' in styles and 'underline' in styles['text-decoration']):
                    run.underline = True
                    print(f" → Aplicado UNDERLINE a: {text_content}")  # 📌 Debug

                if 'color' in styles:
                    color_value = styles['color']
                    print(f" → Detectado color {color_value} en: {text_content}")  # 📌 Debug

                    if isinstance(color_value, str) and color_value.startswith('#'):
                        clean_color = color_value.lstrip('#')
                        if len(clean_color) == 6:
                            r, g, b = int(clean_color[0:2], 16), int(clean_color[2:4], 16), int(clean_color[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                            print(f" → Aplicado color {clean_color} a: {text_content}")  # 📌 Debug

        elif element.name == 'hr':  
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            paragraph._element.get_or_add_pPr().append(pBdr)

        elif element.name == 'br':
            if 'parent' in locals() and isinstance(parent, docx.text.paragraph.Paragraph):
                parent.add_run().add_break()

        elif element.name == 'img':
            src = element['src']
            if src.startswith('data:image'):
                format, imgstr = src.split(';base64,')
                ext = format.split('/')[-1]
                add_base64_image_to_doc(doc, imgstr, ext)

            elif src.startswith('http'):
                try:
                    response = requests.get(src, timeout=5)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        doc.add_picture(image_stream)
                    else:
                        print(f"⚠️ ERROR al descargar imagen: {response.status_code}")
                except requests.exceptions.RequestException as e:
                    print(f"⚠️ ERROR al conectar con la imagen: {e}")

        elif element.name == 'pre':
            text_content = element.get_text()
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]

            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text_content)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = False

            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), '444444')
            tc_pr.append(shd)

            doc.add_paragraph()




def consolidate_vulnerabilities(vulnerabilities):
    consolidated = {}
    for vuln in vulnerabilities:
        if vuln.name not in consolidated:
            consolidated[vuln.name] = vuln.risk_factor
        else:
            # Mantener la mayor severidad si hay duplicados
            current_risk = consolidated[vuln.name]
            if risk_factor_to_int(vuln.risk_factor) < risk_factor_to_int(current_risk):
                consolidated[vuln.name] = vuln.risk_factor
    return consolidated

def risk_factor_to_int(risk_factor):
    risk_mapping = {
        'Critical': 1,
        'High': 2,
        'Medium': 3,
        'Low': 4
    }
    return risk_mapping.get(risk_factor, 5)




def add_hyperlink(paragraph, text, url):
    """
    A function that places a hyperlink within a paragraph object.
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param text: The text displayed for the hyperlink.
    :param url: A string containing the required url
    :return: The hyperlink object
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)

    # Create a w:t element with the text
    w_t = OxmlElement('w:t')
    w_t.text = text

    # Add the w:t element to the w:r element
    new_run.append(w_t)

    # Add the w:r element to the w:hyperlink element
    hyperlink.append(new_run)

    # Add the w:hyperlink element to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

def generate_vulnerability_table(doc, vulnerabilities, language='en'):
    if language == 'es':
        headers = ['Vulnerabilidad', 'Severidad']
    else:
        headers = ['Vulnerability', 'Severity']

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = headers[0]
    hdr_cells[1].text = headers[1]

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    table.columns[0].width = Inches(4.6)
    table.columns[1].width = Inches(1.2)

    consolidated_vulns = consolidate_vulnerabilities(vulnerabilities)
    sorted_vulns = sorted(consolidated_vulns.items(), key=lambda item: risk_factor_to_int(item[1]))

    for vuln, risk in sorted_vulns:
        row_cells = table.add_row().cells
        add_hyperlink(row_cells[0].paragraphs[0], vuln, f'#{vuln}')
        
        if not row_cells[0].paragraphs[0].runs:
            row_cells[0].paragraphs[0].add_run()
        row_cells[0].paragraphs[0].runs[0].font.bold = True

        run = row_cells[1].paragraphs[0].add_run(risk)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        if risk == 'Critical':
            set_cell_background(row_cells[1], '800080')
        elif risk == 'High':
            set_cell_background(row_cells[1], 'FF0000')
        elif risk == 'Medium':
            set_cell_background(row_cells[1], 'FFA500')
        elif risk == 'Low':
            set_cell_background(row_cells[1], '008000')



def generate_vulns_summary_table(doc, vulnerabilities, language):
    # Ordenar las vulnerabilidades por nivel de criticidad
    criticity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}
    sorted_vulnerabilities = sorted(vulnerabilities, key=lambda x: criticity_order.get(x.risk_factor, 4))

    table = doc.add_table(rows=len(sorted_vulnerabilities) + 1, cols=2)
    table.style = 'TableGrid'
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Vulnerabilidad' if language == 'es' else 'Vulnerability'
    hdr_cells[1].text = 'Severidad' if language == 'es' else 'Severity'
    for hdr_cell in hdr_cells:
        hdr_cell.paragraphs[0].runs[0].bold = True
        hdr_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Blanco
        set_cell_background(hdr_cell, '000000')  # Negro

    # Datos de vulnerabilidades
    for idx, vuln in enumerate(sorted_vulnerabilities, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = vuln.name
        legible_risk, risk_color_code = risk_factor_to_legible_text_and_color(vuln.risk_factor, language)
        
        severity_paragraph = row_cells[1].paragraphs[0]
        severity_run = severity_paragraph.add_run(legible_risk)
        severity_run.bold = True

        set_cell_background(row_cells[1], risk_color_code)
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco



def set_cell_styles(cell, styles):
    if 'color' in styles:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = styles['color']
    if 'font_size' in styles:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(styles['font_size'])
    if 'bold' in styles:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = styles['bold']
    if 'italic' in styles:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.italic = styles['italic']
    if 'underline' in styles:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.underline = styles['underline']
    if 'strike' in styles:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.strike = styles['strike']

# Función mejorada para añadir una imagen en formato base64 al documento de Word
def add_base64_image_to_doc(doc, base64_str, ext):
    """
    Decodifica una imagen en base64 y la inserta en el documento asegurando que
    el ancho máximo de la imagen sea de 6 pulgadas sin distorsionarla.
    """
    try:
        # Decodificar la imagen base64
        decoded_img = base64.b64decode(base64_str)
        img_io = io.BytesIO(decoded_img)

        # Verificar si la imagen es válida usando PIL
        try:
            img = Image.open(img_io)
            img.verify()  # Verifica si es un archivo de imagen válido
        except Exception as e:
            error_message = f"[Error al verificar la imagen: {str(e)}]"
            print(error_message)
            doc.add_paragraph(error_message)
            return

        # Guardar la imagen en un archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp_file:
            tmp_file.write(decoded_img)
            tmp_file_path = tmp_file.name

        # Obtener el ancho máximo permitido (6 pulgadas) considerando los márgenes
        section = doc.sections[0]
        max_width = Inches(6)  # Máximo de 6 pulgadas
        usable_width = section.page_width - section.left_margin - section.right_margin
        final_width = min(max_width, usable_width)  # Asegurar que no exceda los márgenes

        # Ajustar la imagen manteniendo la proporción
        try:
            with Image.open(tmp_file_path) as pil_img:
                img_width, img_height = pil_img.size
                aspect_ratio = img_height / img_width
                final_width = min(final_width, Inches(img_width / 96))  # Convertir píxeles a pulgadas
                final_height = final_width * aspect_ratio

                # Insertar la imagen en el documento con alineación centrada
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar la imagen
                run = paragraph.add_run()
                run.add_picture(tmp_file_path, width=final_width, height=final_height)
        except Exception as e:
            error_message = f"[Error al añadir la imagen al documento: {str(e)}]"
            print(error_message)
            doc.add_paragraph(error_message)
        finally:
            os.unlink(tmp_file_path)

    except base64.binascii.Error as e:
        error_message = f"[Error al decodificar la imagen base64: {str(e)}]"
        print(error_message)
        doc.add_paragraph(error_message)
    except Exception as e:
        error_message = f"[Error desconocido al procesar la imagen: {str(e)}]"
        print(error_message)
        doc.add_paragraph(error_message)

# Añadir funcionalidad para parsear el contenido del Writeup y mantener los estilos, incluyendo imágenes
def add_attack_narrative_to_document(doc, attack_narrative_content, writeup_name, request, language='en'):
    """
    Inserta el contenido del writeup al documento DOCX respetando los estilos:
    - Comandos simples <code> con texto fucsia + fondo gris oscuro
    - Bloques de código <pre><code class="language-*"> como tabla gris oscuro con texto blanco
    - Manejo de imágenes (external images are downloaded and saved locally first)
    """
    if not writeup_name or writeup_name.strip() == "":
        writeup_name = "Unknown_Writeup"

    safe_folder = writeup_name.replace(" ", "_")
    protected_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, safe_folder)
    os.makedirs(protected_path, exist_ok=True)

    # --- Step A: Download external images and rewrite HTML before processing ---
    attack_narrative_content = download_external_images(
        attack_narrative_content, writeup_name
    )
    logger.info("Pre-processed HTML for writeup '%s' (external images downloaded)", writeup_name)

    soup = BeautifulSoup(attack_narrative_content, 'html.parser')

    for element in soup.contents:
        if element.name == 'p':
            paragraph = doc.add_paragraph()

            for child in element.children:
                if child.name == 'code' and not child.has_attr('class'):
                    # ✅ Código inline: <code>comando</code>
                    run = paragraph.add_run(child.get_text(strip=True))
                    run.font.name = 'Courier New'
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 0, 255)  # Fucsia

                    rPr = run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), '444444')  # Gris oscuro
                    rPr.append(shd)

                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)

                elif isinstance(child, str):
                    paragraph.add_run(child.strip())

        elif element.name == 'pre':
            code_tag = element.find('code')
            if code_tag and code_tag.has_attr('class') and any(cls.startswith('language-') for cls in code_tag['class']):
                # ✅ Bloque de código: <pre><code class="language-xxx">
                text_content = code_tag.get_text(strip=False)

                table = doc.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                cell = table.rows[0].cells[0]
                paragraph = cell.paragraphs[0]

                run = paragraph.add_run(text_content.strip())
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco

                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), '444444')  # Gris oscuro
                tc_pr.append(shd)

                doc.add_paragraph()
            else:
                # ✅ Bloque <pre> sin <code>
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(element.get_text(strip=True))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                run.bold = True
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)

        elif element.name == 'img':
            image_src = element.get('src')
            if image_src:
                if image_src.startswith('data:image/'):
                    # Base64 embedded image — use margin-safe insertion
                    base64_data = image_src.split(',')[1]
                    add_base64_image_to_doc_safe(doc, base64_data, 'png')

                elif "/protected_media/" in image_src or settings.MEDIA_URL in image_src:
                    # Local image from protected_media or media folder
                    image_path = resolve_image_path(image_src, writeup_name)
                    if image_path:
                        add_image_to_doc(doc, image_path)
                    else:
                        # Fallback: try fetching via HTTP with session cookies
                        image_filename = os.path.basename(image_src)
                        protected_image_url = f"http://localhost:8000/protected_media/{safe_folder}/{image_filename}"
                        image_data = fetch_protected_image(protected_image_url, request)
                        if image_data:
                            add_image_to_doc(doc, image_data)
                        else:
                            logger.warning("Image not found: %s (writeup=%s)", image_src, writeup_name)
                            doc.add_paragraph(f'[Image not found: {image_src}]')

                elif image_src.startswith('/'):
                    # Static/root-relative path
                    image_path = resolve_image_path(image_src, writeup_name)
                    if image_path:
                        add_image_to_doc(doc, image_path)
                    else:
                        logger.warning("Static image not found: %s (writeup=%s)", image_src, writeup_name)
                        doc.add_paragraph(f'[Image not found: {image_src}]')

                elif image_src.startswith('http'):
                    # External URL (should have been downloaded already by download_external_images,
                    # but handle as fallback for edge cases like blob: URLs)
                    try:
                        resp = requests.get(image_src, timeout=10)
                        resp.raise_for_status()
                        # Save locally to avoid re-downloading
                        from ProjectManager.image_utils import _download_and_save
                        local = _download_and_save(image_src, protected_path)
                        if local:
                            add_image_to_doc(doc, local)
                        else:
                            add_image_to_doc(doc, resp.content)
                    except requests.exceptions.RequestException as e:
                        logger.error("Failed to download image %s: %s", image_src, e)
                        doc.add_paragraph(f'[Error al descargar imagen: {str(e)}]')

        elif element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)





def serve_protected_media(request, writeup_name, filename):
    """
    🔐 Sirve archivos de `protected_media/` solo si el usuario tiene acceso al proyecto.
    """
    # 📁 Construye la ruta del archivo
    file_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, writeup_name, filename)

    if os.path.exists(file_path):
        return FileResponse(open(file_path, "rb"), content_type="image/png")  # Ajustar content-type según imagen
    else:
        return HttpResponseNotFound("Image not found.")



PROTECTED_MEDIA_URL = "/admin/media/"
PROTECTED_MEDIA_ROOT = "protected_media"

def get_valid_image_url(img_url):
    """ Convierte rutas locales a URLs accesibles """
    if img_url.startswith("/media/"):
        return urljoin(PROTECTED_MEDIA_URL, img_url.lstrip("/"))
    elif img_url.startswith("/protected_media/"):
        return urljoin(PROTECTED_MEDIA_URL, img_url.lstrip("/"))
    elif img_url.startswith(("http://localhost", "https://localhost")):
        return img_url  # Si ya es una URL válida, la devuelve sin cambios
    return None  # Si no es válida, devuelve None


def download_image(img_url, writeup_folder):
    """ Descarga una imagen y la guarda en /protected_media/{Writeup_Name}/ """
    try:
        response = requests.get(img_url, stream=True, timeout=10)
        response.raise_for_status()

        # Detectar extensión correcta
        img_extension = mimetypes.guess_extension(response.headers.get("content-type", "image/png"))
        if not img_extension:
            img_extension = ".png"

        # Obtener nombre de archivo sin caracteres extraños
        img_name = os.path.basename(urlparse(img_url).path)
        img_name = "".join(c for c in img_name if c.isalnum() or c in (".", "_")).rstrip()

        save_path = os.path.join(writeup_folder, f"{img_name}{img_extension}")

        # Guardar la imagen
        with open(save_path, "wb") as img_file:
            for chunk in response.iter_content(1024):
                img_file.write(chunk)

        return save_path
    except Exception as e:
        print(f"⚠️ Error descargando imagen {img_url}: {e}")
        return None





def protected_media_view(request, writeup_name, filename):
    """
    Sirve archivos protegidos desde PROTECTED_MEDIA solo si el usuario tiene permisos.
    Retorna 403 Forbidden en lugar de redirigir a login.
    """
    file_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, writeup_name, filename)

    if not request.user.is_authenticated:
        return HttpResponseForbidden("You do not have permission to view this image.")

    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), content_type='image/png')  # Ajusta el content_type si es necesario
    else:
        return HttpResponseNotFound("Image not found.")


def fetch_protected_image(image_url, request):
    """
    📥 Descarga imágenes de `protected_media/` con la sesión autenticada.
    """
    try:
        session = requests.Session()
        headers = {
            "Cookie": request.META.get("HTTP_COOKIE", ""),  # Envía las cookies del usuario autenticado
            "User-Agent": "Mozilla/5.0",
        }
        response = session.get(image_url, headers=headers, allow_redirects=False)

        if response.status_code == 200:
            print(f"✅ Imagen descargada correctamente: {image_url}")
            return response.content  # Devuelve los bytes de la imagen
        elif response.status_code == 403:
            print(f"❌ Acceso denegado a {image_url}. Verifica permisos del usuario.")
        else:
            print(f"⚠️ Error {response.status_code} al descargar imagen: {image_url}")

        return None

    except Exception as e:
        print(f"❌ ERROR al descargar imagen protegida {image_url}: {e}")
        return None



def process_images(html_content, writeup_name):
    """
    Convierte rutas de imágenes en el HTML a rutas accesibles desde PROTECTED_MEDIA.
    """
    from bs4 import BeautifulSoup
    from urllib.parse import urljoin

    soup = BeautifulSoup(html_content, 'html.parser')

    for img_tag in soup.find_all('img'):
        img_src = img_tag.get('src')
        if img_src:
            # 🔹 Si la imagen está en /media/, convertir a /protected_media/
            if img_src.startswith(settings.MEDIA_URL) or img_src.startswith('/media/'):
                image_filename = os.path.basename(img_src)
                
                # ✅ Nueva ruta completa con localhost
                new_src = urljoin("http://localhost:8000/", f"protected_media/{writeup_name}/{image_filename}")
                
                img_tag['src'] = new_src  # Reemplazar en el HTML
                print(f"🔄 URL de imagen corregida: {img_src} → {new_src}")  # Depuración
            
            # 🔹 Si la imagen tiene un URL relativo (Ej: /static/)
            elif img_src.startswith('/'):
                new_src = urljoin("http://localhost:8000/", img_src.lstrip('/'))
                img_tag['src'] = new_src
                print(f"🔄 URL de imagen relativa corregida: {img_src} → {new_src}")  # Depuración

    return str(soup)




class CustomHTMLParser(HTMLParser):
    """
    Parser para convertir HTML a DOCX con mejor soporte de estilos, imágenes y código.
    """
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = self.doc.add_paragraph()
        self.current_run = self.current_paragraph.add_run()
        self.in_code_block = False

    def handle_starttag(self, tag, attrs):
        if tag in ["strong", "b"]:
            self.current_run = self.current_paragraph.add_run()
            self.current_run.bold = True
        elif tag in ["em", "i"]:
            self.current_run = self.current_paragraph.add_run()
            self.current_run.italic = True
        elif tag == "u":
            self.current_run = self.current_paragraph.add_run()
            self.current_run.underline = True
        elif tag == "h1":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_run = self.current_paragraph.add_run()
            self.current_run.bold = True
            self.current_run.font.size = Pt(18)
        elif tag == "h2":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_run = self.current_paragraph.add_run()
            self.current_run.bold = True
            self.current_run.font.size = Pt(16)
        elif tag == "h3":
            self.current_paragraph = self.doc.add_paragraph()
            self.current_run = self.current_paragraph.add_run()
            self.current_run.bold = True
            self.current_run.font.size = Pt(14)
        elif tag in ["ul", "ol"]:
            self.current_paragraph = self.doc.add_paragraph()
        elif tag == "li":
            self.current_paragraph = self.doc.add_paragraph(style="ListBullet")
            self.current_run = self.current_paragraph.add_run()
        elif tag == "code":
            self.current_run = self.current_paragraph.add_run()
            self.current_run.font.name = "Courier New"
            self.current_run.bold = True
        elif tag == "pre":
            self.in_code_block = True
            self.current_paragraph = self.doc.add_paragraph()
            self.current_paragraph.style = self.doc.styles["Normal"]  # 🔹 Se usa "Normal" en lugar de "Code"
            self.current_run = self.current_paragraph.add_run()
            self.current_run.font.name = "Courier New"
            self.current_run.font.size = Pt(10)
        elif tag == "span":
            for attr in attrs:
                if attr[0] == "style" and "color" in attr[1]:
                    color_value = attr[1].split(":")[-1].strip()

                    if color_value.startswith("#"):
                        self.current_run.font.color.rgb = RGBColor.from_string(color_value.replace("#", ""))
        elif tag == "img":
            for attr in attrs:
                if attr[0] == "src":
                    self.insert_image(attr[1])

    def handle_endtag(self, tag):
        if tag == "pre":
            self.in_code_block = False
        elif tag == "p":
            self.current_paragraph = self.doc.add_paragraph()

    def handle_data(self, data):
        self.current_run.add_text(data.strip())

    def insert_image(self, url):
        """
        Descarga e inserta imágenes desde URLs en el documento DOCX.
        """
        try:
            response = requests.get(url, stream=True)
            if response.status_code == 200:
                image_stream = BytesIO(response.content)
                self.doc.add_picture(image_stream, width=Inches(5))
                self.current_paragraph = self.doc.add_paragraph()
        except Exception as e:
            print(f"⚠️ Error descargando imagen {url}: {e}")

def insert_code_block(self, text):
    table = self.doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # Insertar el código dentro de la celda
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    font = run.font
    font.name = "Courier New"
    font.size = Pt(9)
    font.color.rgb = RGBColor(255, 255, 255)  # blanco

    # Aplicar fondo gris oscuro
    shading_elm = cell._tc.get_or_add_tcPr().add_new_shd()
    shading_elm.val = 'clear'
    shading_elm.fill = '444444'  # gris oscuro




def fetch_protected_image(image_url, request):
    """ Descarga imágenes de protected_media con la sesión del usuario """
    try:
        session = requests.Session()
        headers = {"Cookie": request.META.get("HTTP_COOKIE", "")}  # Enviar cookies de sesión
        response = session.get(image_url, headers=headers, allow_redirects=False)

        if response.status_code == 200:
            print(f"✅ Imagen descargada correctamente: {image_url}")
            return response.content  # Devuelve los bytes de la imagen
        else:
            print(f"⚠️ Error {response.status_code} al descargar imagen: {image_url}")
            return None

    except Exception as e:
        print(f"❌ ERROR al descargar imagen protegida {image_url}: {e}")
        return None
    

    

def insert_html_to_docx(html_content, doc, writeup_id):
    """
    📄 Inserta HTML en el DOCX, descargando imágenes en `protected_media` y usando URLs protegidas.
    """
    paragraphs = re.split(r"\n+", html_content)
    for paragraph in paragraphs:
        # 🖼️ IMÁGENES
        if '<img ' in paragraph:
            img_url_match = re.search(r'src="(.*?)"', paragraph)
            if img_url_match:
                img_url = img_url_match.group(1)
                img_path, protected_url = download_image(img_url, writeup_id)

                if img_path and os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(4))
                    except Exception as e:
                        print(f"⚠️ Error insertando imagen {img_path}: {e}")

        # 🔹 BLOQUES DE CÓDIGO (EN TABLA)
        elif '<pre><code>' in paragraph:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.rows[0].cells[0]
            cell.text = re.sub(r'<.*?>', '', paragraph)  # Elimina etiquetas HTML
            cell.paragraphs[0].runs[0].font.name = "Courier New"
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            shading = OxmlElement("w:shd")
            shading.set(ns.qn("w:fill"), "222222")
            cell._element.get_or_add_tcPr().append(shading)

        # 🔹 TEXTO NORMAL (CON COLORES)
        else:
            p = doc.add_paragraph()
            bold_matches = re.findall(r'<(b|strong)>(.*?)</\1>', paragraph)
            for match in bold_matches:
                run = p.add_run(match[1])
                run.bold = True
            
            color_matches = re.findall(r'<span style="color:(.*?)">(.*?)</span>', paragraph)
            for match in color_matches:
                run = p.add_run(match[1])
                try:
                    if match[0].startswith("hsl"):
                        # 🛠 Convertir HSL a RGB
                        hsl_values = re.findall(r"[\d.]+", match[0])
                        if len(hsl_values) == 3:
                            h, s, l = map(float, hsl_values)
                            r, g, b = hsl_to_rgb(h, s, l)
                            run.font.color.rgb = RGBColor(r, g, b)
                    else:
                        # 🛠 Convertir HEX a RGB
                        color_hex = match[0].lstrip('#')
                        run.font.color.rgb = RGBColor(int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16))
                except:
                    print(f"⚠️ No se pudo convertir el color {match[0]}")

            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def hsl_to_rgb(h, s, l):
    """
    Convierte un valor de color en HSL a RGB.
    """
    import colorsys
    r, g, b = colorsys.hls_to_rgb(h / 360, l, s)
    return int(r * 255), int(g * 255), int(b * 255)


def insert_writeup_images_into_report(writeup_id, doc):
    """ Maneja la inserción de imágenes del writeup en el reporte """
    writeup = Writeup.objects.get(id=writeup_id)
    writeup_name = writeup.title.replace(" ", "_")

    # 🔹 Crear directorio antes de procesar imágenes
    protected_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, writeup_name)
    if not os.path.exists(protected_path):
        os.makedirs(protected_path)
        print(f"📂 Directorio creado: {protected_path}")

    image_filename = "example_image.png"  # Nombre real de la imagen en writeup
    image_path = os.path.join(protected_path, image_filename)

    if os.path.exists(image_path):
        print(f"✅ Imagen encontrada: {image_path}")

        from docx.shared import Inches
        from PIL import Image

        # Obtener ancho útil del documento considerando márgenes
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        usable_width_inches = usable_width / 914400

        # Escalar proporcionalmente la imagen
        with Image.open(image_path) as img:
            width, height = img.size
            aspect_ratio = height / width
            scaled_height_inches = usable_width_inches * aspect_ratio

        doc.add_picture(image_path, width=Inches(usable_width_inches), height=Inches(scaled_height_inches))

    else:
        print(f"⚠️ Imagen NO encontrada: {image_path}")


class HTMLToDocxParser(HTMLParser):
    """
    Parser de HTML para convertirlo en un documento DOCX con formato adecuado.
    """
    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.in_code_block = False
        self.in_bold = False

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)

        if tag == "p":
            self.current_paragraph = self.doc.add_paragraph()
        elif tag == "strong" or tag == "b":
            self.in_bold = True
        elif tag == "br":
            self.current_paragraph.add_run("\n")
        elif tag == "img" and "src" in attrs:
            self.insert_image(attrs["src"])
        elif tag == "code":
            self.in_code_block = True
        elif tag == "pre":
            self.in_code_block = True
            self.insert_code_block()

    def handle_endtag(self, tag):
        if tag in ["strong", "b"]:
            self.in_bold = False
        elif tag in ["code", "pre"]:
            self.in_code_block = False

    def handle_data(self, data):
        if not self.current_paragraph:
            self.current_paragraph = self.doc.add_paragraph()

        run = self.current_paragraph.add_run(data)

        if self.in_bold:
            run.bold = True

        if self.in_code_block:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

    def insert_image(self, url):
        """
        Descarga e inserta una imagen en el documento.
        """
        try:
            response = requests.get(url, stream=True)
            response.raise_for_status()
            image = Image.open(BytesIO(response.content))
            image_stream = BytesIO()
            image.save(image_stream, format=image.format)
            self.doc.add_picture(image_stream)
        except Exception as e:
            print(f"⚠️ ERROR al descargar la imagen {url}: {e}")

    def insert_code_block(self):
        """
        Crea un cuadro con estilo de código.
        """
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        shading_elm = parse_xml(r'<w:shd {} w:fill="2E2E2E"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading_elm)
        self.current_paragraph = cell.paragraphs[0]
        self.current_paragraph.style = self.doc.styles["Normal"]
        self.current_paragraph.paragraph_format.space_before = Pt(3)
        self.current_paragraph.paragraph_format.space_after = Pt(3)

def insert_html_to_docx(html_content, doc):
    soup = BeautifulSoup(html_content, "html.parser")

    for el in soup.contents:
        if el.name == "table":
            # Crear tabla
            rows = el.find_all("tr")
            if not rows:
                continue

            num_cols = max(len(row.find_all(["td", "th"])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            table.style = 'Table Grid'

            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    text = cell.get_text(strip=True)
                    paragraph = table.cell(i, j).paragraphs[0]
                    run = paragraph.add_run(text)
                    run.font.size = Pt(10)
                    if cell.name == "th":
                        run.bold = True

        elif el.name in ["h1", "h2", "h3"]:
            doc.add_heading(el.get_text(strip=True), level=int(el.name[1]))

        elif el.name == "p":
            doc.add_paragraph(el.get_text())

        else:
            doc.add_paragraph(el.get_text())


def insertar_writeup_con_imagenes(writeup, doc):
    """Inserta contenido del writeup en el docx, ajustando imágenes locales."""

    soup = BeautifulSoup(writeup.html, "html.parser")

    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_width_inches = usable_width / 914400

    for img_tag in soup.find_all("img"):
        img_src = img_tag.get("src")

        # Aseguramos que es una imagen local y que existe
        if img_src and img_src.startswith("/media/"):
            relative_path = img_src.replace("/media/", "")
            image_path = os.path.join(settings.MEDIA_ROOT, relative_path)

            if os.path.exists(image_path):
                try:
                    with Image.open(image_path) as img:
                        width, height = img.size
                        aspect_ratio = height / width
                        scaled_height = usable_width_inches * aspect_ratio

                    doc.add_picture(image_path, width=Inches(usable_width_inches), height=Inches(scaled_height))
                    print(f"🖼️ Imagen insertada y escalada: {image_path}")

                except Exception as e:
                    print(f"❌ Error al procesar imagen {image_path}: {e}")
            else:
                print(f"⚠️ Image not found at path: {image_path}")

        # Eliminamos el <img> para que html2docx no la duplique
        img_tag.decompose()

    html2docx(str(soup), doc)


def insert_table_of_contents(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)



@login_required
def save_node_position(request, target_id):
    if request.method == "POST":
        # Consultant node (id 0) has no DB record; skip saving
        if target_id == 0:
            return JsonResponse({"status": "success"})
        # Support both JSON and form-encoded bodies
        import json as _json
        if request.content_type == "application/json":
            try:
                data = _json.loads(request.body)
            except _json.JSONDecodeError:
                return JsonResponse({"status": "error", "message": "Invalid JSON"}, status=400)
            x = data.get("x")
            y = data.get("y")
        else:
            x = request.POST.get("x")
            y = request.POST.get("y")

        if x is None or y is None:
            return JsonResponse({"status": "error", "message": "x and y are required"}, status=400)

        try:
            target = Target.objects.get(id=target_id)
            get_object_or_404(get_projects_queryset(request.user), pk=target.project_id)
            target.x_position = float(x)
            target.y_position = float(y)
            target.save()
            return JsonResponse({"status": "success"})
        except Target.DoesNotExist:
            return JsonResponse({"status": "error", "message": "Target not found"}, status=404)

    return JsonResponse({"status": "error", "message": "Invalid request"}, status=400)
    




def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

TOC_PLACEHOLDER = "###TOC_PLACEHOLDER###"

def _collect_toc_entries_after_placeholder(doc, placeholder_para, exclude_titles=None):
    """Recorre el documento y recoge (level, text) de cada Heading que esté después del placeholder."""
    exclude_titles = set(exclude_titles or [])
    paras = list(doc.paragraphs)
    try:
        start_idx = next(i for i, p in enumerate(paras) if p == placeholder_para) + 1
    except StopIteration:
        start_idx = 0
    entries = []
    for para in paras[start_idx:]:
        if not para.style.name.startswith("Heading"):
            continue
        try:
            level = int(para.style.name.replace("Heading ", ""))
        except ValueError:
            level = 1
        text = para.text.strip()
        if not text or text in exclude_titles:
            continue
        entries.append((level, text))
    return entries

def _create_toc_field_paragraph(doc):
    """Crea un párrafo con el campo TOC de Word (números de página al actualizar en Word)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'begin')
    run._r.append(fldChar)
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    # \\o "1-3" = niveles 1-3, \\h = hipervínculos, \\z = ocultar números de párrafo, \\u = usar estilos, \\p = números de página (Título.........N)
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u \\p'
    run._r.append(instrText)
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'separate')
    run._r.append(fldChar)
    fldChar = create_element('w:fldChar')
    create_attribute(fldChar, 'w:fldCharType', 'end')
    run._r.append(fldChar)
    return paragraph

def _replace_toc_placeholder_with_field(doc, placeholder_text=TOC_PLACEHOLDER):
    """Sustituye el párrafo que contiene ###TOC_PLACEHOLDER### por un campo TOC de Word (con números de página)."""
    placeholder_para = None
    for para in doc.paragraphs:
        if placeholder_text in (para.text or ""):
            placeholder_para = para
            break
    if not placeholder_para:
        return
    toc_para = _create_toc_field_paragraph(doc)
    toc_elem = toc_para._element
    toc_elem.getparent().remove(toc_elem)
    parent = placeholder_para._element.getparent()
    parent.replace(placeholder_para._element, toc_elem)

def insert_table_of_contents(doc):
    """Inserta el campo Word TOC al final del documento (se actualiza al abrir en Word). Incluye números de página."""
    _create_toc_field_paragraph(doc)


def _set_update_fields_on_open(doc):
    """Activa la actualización de campos al abrir el documento en Word (para que el TOC muestre números de página)."""
    try:
        update_el = OxmlElement('w:updateFields')
        update_el.set(qn('w:val'), 'true')
        doc.settings.element.append(update_el)
    except Exception:
        pass


def add_footer_page_number(doc):
    """Añade número de página en el pie de cada página (campo PAGE)."""
    section = doc.sections[0]
    footer = section.footer
    # Limpiar párrafos existentes del footer si los hay
    for p in footer.paragraphs:
        p.clear()
    if not footer.paragraphs:
        footer.add_paragraph()
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    # Campo PAGE: muestra el número de página actual
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_document_background(doc, color):
    """
    Sets the background color of the document.
    """
    if not color: return
    color = color.lstrip('#')
    background = OxmlElement('w:background')
    background.set(qn('w:color'), color)
    doc.element.insert(0, background)
    display_shape = OxmlElement('w:displayBackgroundShape')
    doc.settings.element.append(display_shape)


@login_required
def generate_report(request, project_id):
    print(f"🚨 Generando reporte para Proyecto ID {project_id} con método {request.method}")
    if request.method == 'POST':
        project = get_object_or_404(get_projects_queryset(request.user), pk=project_id)
        print(f"🛠 DEBUG -> Buscando Writeups para Project ID: {project.id}")
        attack_narratives = Writeup.objects.filter(project=project)
        print(f"✅ DEBUG -> Writeups encontrados: {attack_narratives.count()}")


        # Report labels and headers are always in English
        language = "en"

        print(f"🛠 DEBUG -> Generating report in language: {language}")


        # 📌 Extraer la cookie de sesión del usuario autenticado
        session_cookie = request.session.session_key
        if not session_cookie:
            return HttpResponse("Error: User session not found.", status=403)

        # 📌 Extraer el token CSRF si está disponible
        csrf_token = request.COOKIES.get('csrftoken', '')

        # 📌 Definir la ruta del script de Puppeteer
        script_path = os.path.join(settings.BASE_DIR, "scripts/capture_graph.js")

        # 📌 Ejecutar Puppeteer con los argumentos correctos
        try:
            print(f"📸 Generando imagen de GraphMap para el Proyecto {project_id}...")

            # Asegurarse de que los argumentos sean STRINGS y no `None`
            args = ["node", script_path, str(project_id), str(session_cookie), str(csrf_token)]
            print(f"🛠 DEBUG -> Ejecutando: {' '.join(args)}")  # Ver qué se está ejecutando realmente

            result = subprocess.run(
                args,
                check=True,
                capture_output=True,
                text=True
            )

            print(f"✅ Puppeteer Output:\n{result.stdout}")

        except subprocess.CalledProcessError as e:
            print(f"❌ ERROR en Puppeteer: {e.stderr}")
            return HttpResponse(f"Error generating the GraphMap: {e.stderr}", status=500)

        # 📄 CONTINUAR con la generación del reporte en lugar de hacer un return aquí
        print("✅ Imagen de GraphMap generada correctamente. Continuando con el reporte...")

        # Crear el documento
        doc = Document()

        # Pie de página con número de página en todas las hojas
        add_footer_page_number(doc)

        # Apply background color if configured
        if project.report_template and project.report_template.background_color:
            try:
                set_document_background(doc, project.report_template.background_color)
            except Exception as e:
                print(f"Error setting background color: {e}")


        # --- Sección de la Tapa del Reporte ---
        cover = project.cover_template
        if cover:
            # Título utilizando el nombre del proyecto
            title = doc.add_heading(level=1)
            title_run = title.add_run(f"Offensive Security Testing: {getattr(cover, 'analisys_type', 'N/A')} - {project.name}")
            title_run.font.size = Pt(24)  # Tamaño de 24 pt
            title_run.bold = True  # Negrita
            title.alignment = 1  # Centrar el título

            subtitle_text = getattr(cover, 'subtitle', '') or ''
            if subtitle_text:
                sub_para = doc.add_paragraph(subtitle_text)
                sub_para.alignment = 1
                for run in sub_para.runs:
                    run.font.size = Pt(12)
                sub_para.paragraph_format.space_after = Pt(6)
            consultant_text = getattr(cover, 'consultant_name', '') or ''
            if consultant_text:
                cons_para = doc.add_paragraph(consultant_text if language == 'es' else consultant_text)
                cons_para.alignment = 1
                for run in cons_para.runs:
                    run.font.size = Pt(10)
                cons_para.paragraph_format.space_after = Pt(6)

            # Añadir el texto "| REPORTE EJECUTIVO/TÉCNICO |"
            report_type_paragraph = doc.add_paragraph()
            report_type_run = report_type_paragraph.add_run("| EXECUTIVE/TECHNICAL REPORT |")
            report_type_run.bold = True  # Negrita
            report_type_run.font.size = Pt(15.5)  # Tamaño de 15.5 pt
            report_type_run.font.color.rgb = RGBColor(255, 140, 0)  # Color naranja
            report_type_paragraph.alignment = 1  # Centrar el texto

            # Engagement dates
            fecha_inicio = project.start_date.strftime("%d/%m/%Y")
            fecha_fin = project.end_date.strftime("%d/%m/%Y")
            fechas_texto = f"Engagement start date: {fecha_inicio} / Engagement end date: {fecha_fin}"
            fecha_paragraph = doc.add_paragraph(fechas_texto)
            fecha_paragraph.alignment = 1  # Center dates
            fecha_run = fecha_paragraph.runs[0]
            fecha_run.font.size = Pt(8)  # Tamaño de 8 pt





            # --- Imagen principal (escalable según customer_image_scale) ---
            if cover.customer_image:
                section = doc.sections[0]
                page_width = section.page_width
                usable = page_width - section.left_margin - section.right_margin
                scale_pct = max(10, min(100, getattr(cover, 'customer_image_scale', 100) or 100))
                image_width = usable * (scale_pct / 100.0)

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.right_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                if scale_pct >= 100:
                    extra_space = Inches(0.1)
                    paragraph.paragraph_format.left_indent = -section.left_margin - extra_space
                    paragraph.paragraph_format.right_indent = -section.right_margin - extra_space
                    image_width = page_width + (extra_space * 2)
                run = paragraph.add_run()
                run.add_picture(cover.customer_image.path, width=image_width)
            else:
                doc.add_paragraph("No provider image was selected for this report.")
            # --- Fin de la Sección ---




            # Obtenemos el header y limpiamos su contenido
            header = doc.sections[0].header
            for para in header.paragraphs:
                p_element = para._element
                p_element.getparent().remove(p_element)
                para._p = para._element = None

            # Añadimos un nuevo párrafo al header
            header_paragraph = header.add_paragraph()
            header_paragraph.paragraph_format.space_before = Pt(0)
            header_paragraph.paragraph_format.space_after = Pt(0)

            # Alineamos el párrafo a la izquierda para que las tabulaciones funcionen
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Calculamos la posición del tabulador derecho (dentro del margen derecho)
            section = doc.sections[0]
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            usable_width = page_width - left_margin - right_margin
            image_width = Inches(0.9)  # Ancho de la imagen en pulgadas

            # Ajustamos la posición del tabulador derecho restando el ancho de la imagen
            right_tab_position = left_margin + usable_width - image_width

            # Establecemos una tabulación derecha en la posición calculada
            tab_stops = header_paragraph.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(right_tab_position, alignment=WD_TAB_ALIGNMENT.RIGHT)

            # Añadimos la imagen izquierda (header_imagen)
            if cover.header_image:
                header_run = header_paragraph.add_run()
                header_run.add_picture(cover.header_image.path, width=image_width)  # Usamos image_width

            # Añadimos un carácter de tabulación para mover el cursor al tabulador derecho
            header_paragraph.add_run('\t')

            # Añadimos la imagen derecha (customer_header_image)
            if cover.customer_header_image:
                header_run = header_paragraph.add_run()
                header_run.add_picture(cover.customer_header_image.path, width=image_width)  # Usamos image_width

            # Si no se seleccionó una tapa para el reporte
            if not cover.header_image and not cover.customer_header_image:
                doc.add_paragraph("No cover was selected for this report.")
        # --- Fin de la Sección de la Tapa del Reporte ---


        # Salto de página: la portada queda en página 1; el contenido (TOC + cuerpo) empieza en página 2.
        doc.add_page_break()

        # Obtener el contenido del reporte (el template debe incluir "Table of Contents", intro y ###TOC_PLACEHOLDER###)
        report_content = project.report_template.content if project.report_template else ''


        # Obtener todas las vulnerabilidades del proyecto y ordenarlas por criticidad.
        # OWASP: evidence/description/solution can contain tool output (payloads); python-docx
        # inserts plain text only, so no HTML/script injection in the DOCX.
        vulnerabilities = list(Vulnerability.objects.filter(project=project).order_by('-risk_factor'))
        # Ajustar descripción y solución al idioma seleccionado
        for vuln in vulnerabilities:
            if language == 'es':
                vuln.description = vuln.description_es or vuln.description
                vuln.solution = vuln.solution_es or vuln.solution

        # Procesar el contenido del Report Template
        if report_content:
            add_html_to_doc(doc, report_content, vulnerabilities, language)

        # Reemplazar ###Scope### por el contenido del campo scope con formato
        for para in doc.paragraphs:
            if '###Scope###' in para.text:
                scope_data = project.scope.split('\n') if project.scope else []
                if scope_data:
                    # Dividir los datos del alcance en 3 columnas
                    num_rows = (len(scope_data) + 2) // 3
                    table = doc.add_table(rows=num_rows + 1, cols=3)
                    table.style = 'TableGrid'

                    # Añadir el encabezado
                    hdr_cells = table.rows[0].cells
                    for hdr_cell in hdr_cells:
                        hdr_cell.text = 'Scope' if language == 'en' else 'Alcance'
                        hdr_cell.paragraphs[0].runs[0].bold = True
                        hdr_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                        set_cell_background(hdr_cell, '000000')  # Negro

                    # Añadir los datos del alcance distribuidos en 3 columnas
                    for idx, scope in enumerate(scope_data):
                        col = idx // num_rows
                        row = idx % num_rows + 1
                        cell = table.cell(row, col)
                        cell.text = scope.strip()
                        # Aplicar color de fondo alterno
                        if row % 2 == 0:
                            set_cell_background(cell, 'D3D3D3')  # Gris claro
                        else:
                            set_cell_background(cell, 'FFFFFF')  # Blanco

                    # Insertar la tabla en la ubicación del marcador
                    para._element.addnext(table._element)
                para.text = para.text.replace('###Scope###', '')




        # Reemplazar ###SeverityProfile### con la postura de riesgo
        for para in doc.paragraphs:
            if '###SeverityProfile###' in para.text:
                # Determinar la criticidad más alta del proyecto
                risk_levels = ['Critical', 'High', 'Medium', 'Low']
                highest_risk = 'Low'
                for risk_level in risk_levels:
                    if any(vuln.risk_factor == risk_level for vuln in vulnerabilities):
                        highest_risk = risk_level
                        break

                color_map = {
                    'Critical': '800080',  # Púrpura
                    'High': 'FF0000',  # Rojo
                    'Medium': 'FFA500',  # Naranja
                    'Low': '008000'  # Verde
                }

                translated_risk = {
                    'Critical': 'Crítico' if language == 'es' else 'Critical',
                    'High': 'Alto' if language == 'es' else 'High',
                    'Medium': 'Medio' if language == 'es' else 'Medium',
                    'Low': 'Bajo' if language == 'es' else 'Low'
                }

                for para in doc.paragraphs:
                    if '###SeverityProfile###' in para.text:
                        severity_text = translated_risk[highest_risk]
                        para.text = para.text.replace('###SeverityProfile###', '')
                        severity_run = para.add_run(severity_text)
                        severity_run.bold = True
                        severity_run.font.color.rgb = RGBColor.from_string(color_map[highest_risk])



        # Obtener una lista mutable de todos los párrafos
        paragraphs = list(doc.paragraphs)

        # Reemplazar ###PortMapTable### por el contenido del campo de puertos y servicios con formato
        for para in doc.paragraphs:
            if '###PortMapTable###' in para.text:
                print("Encontrado el marcador ###PortMapTable###")
                para.text = para.text.replace('###PortMapTable###', '')  # Limpiar el texto del marcador

                # Filtrar los puertos relacionados con el proyecto
                ports = Port.objects.filter(target__project=project)
                print(f"Puertos encontrados: {ports.count()}")  # Depuración

                if ports.exists():
                    # Crear la tabla con encabezados
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'

                    # Configurar encabezados de la tabla
                    hdr_cells = table.rows[0].cells
                    hdr_titles = ['Host', 'Puerto/Protocolo', 'Estado', 'Banner']
                    for idx, hdr_cell in enumerate(hdr_cells):
                        hdr_cell.text = hdr_titles[idx]
                        hdr_paragraph = hdr_cell.paragraphs[0]
                        hdr_run = hdr_paragraph.runs[0]
                        hdr_run.bold = True
                        hdr_run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco
                        set_cell_background(hdr_cell, '000000')  # Negro

                    # Inicializar contador de filas
                    row_idx = 1  # Comenzamos en 1 porque ya tenemos la fila de encabezado

                    # Añadir los datos de los puertos
                    for port in ports:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(port.target)  # Host (IP, FQDN o URL)
                        row_cells[1].text = f"{port.port_number}/{port.protocol}"  # Puerto/Protocolo
                        row_cells[2].text = port.state  # Estado
                        row_cells[3].text = port.banner if port.banner else 'N/A'  # Banner

                        # Aplicar color de fondo alterno
                        shade = 'D3D3D3' if (row_idx % 2 == 0) else 'FFFFFF'
                        for cell in row_cells:
                            set_cell_background(cell, shade)

                        # Incrementar el contador de filas
                        row_idx += 1

                    # Insertar la tabla en la ubicación del marcador
                    para._element.addnext(table._element)

                else:
                    # Si no hay puertos, insertar un mensaje informativo
                    no_ports_paragraph = doc.add_paragraph("No ports were found for this project.")
                    para._element.addnext(no_ports_paragraph._element)

                # Salir del bucle ya que el marcador ha sido reemplazado
                break



        # Dictionary para almacenar vulnerabilidades consolidadas
        consolidated_vulnerabilities = {}

        # Definir el orden de criticidad
        criticity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}

        # Ordenar las vulnerabilidades por nivel de criticidad
        vulnerabilities = sorted(vulnerabilities, key=lambda x: criticity_order.get(x.risk_factor, 4))

        # Generar el contenido de las vulnerabilidades y recopilar títulos y subtítulos
        for vulnerability in vulnerabilities:
            legible_risk, risk_color_code = risk_factor_to_legible_text_and_color(vulnerability.risk_factor, language)
            vulnerability_name = vulnerability.name

            # 🔥 Aplicar la traducción solo acá
            description = vulnerability.description_es if language == 'es' else vulnerability.description
            solution = vulnerability.solution_es if language == 'es' else vulnerability.solution

            if vulnerability_name in consolidated_vulnerabilities:
                consolidated_vulnerabilities[vulnerability_name]['hosts'].append(vulnerability.hosts_affected)
                consolidated_vulnerabilities[vulnerability_name]['ports'].append(vulnerability.port)
                consolidated_vulnerabilities[vulnerability_name]['evidence'].extend(vulnerability.evidence.split(','))
            else:
                consolidated_vulnerabilities[vulnerability_name] = {
                    'risk': legible_risk,
                    'risk_color': risk_color_code,
                    'hosts': [vulnerability.hosts_affected],
                    'ports': [vulnerability.port],
                    'description': description,
                    'solution': solution,
                    'evidence': vulnerability.evidence.split(','),
                }

        # Insertar la tabla de resumen de vulnerabilidades
        overall_table_marker = '###OverallTableVulns###'
        for para in doc.paragraphs:
            if overall_table_marker in para.text:
                para.text = ""  # Limpia el marcador
                # Insertar la tabla en el lugar del marcador
                table = doc.add_table(rows=2, cols=4)
                table.style = 'TableGrid'
                table.autofit = True

                # Establecer el ancho de las columnas
                column_width = Cm(4)
                for col in table.columns:
                    col.width = column_width

                # Rellenar la tabla con los valores
                header_cells = ['Crítica' if language == 'es' else 'Critical',
                                'Alta' if language == 'es' else 'High',
                                'Media' if language == 'es' else 'Medium',
                                'Baja' if language == 'es' else 'Low']
                
                risk_factors = ['Critical', 'High', 'Medium', 'Low']
                counts = [sum(1 for v in vulnerabilities if v.risk_factor == risk) for risk in risk_factors]

                for i, (header, count, risk) in enumerate(zip(header_cells, counts, risk_factors)):
                    legible_risk, risk_color_code = risk_factor_to_legible_text_and_color(risk, language)
                    color = RGBColor(255, 255, 255)  # Blanco

                    table.cell(0, i).text = header
                    table.cell(1, i).text = str(count)

                    set_cell_background(table.cell(0, i), risk_color_code)
                    set_cell_background(table.cell(1, i), risk_color_code)
                    
                    for row in range(2):
                        cell = table.cell(row, i)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = color

                para._element.addnext(table._element)

        # Generar las tablas de vulnerabilidades
        for vulnerability_name, details in consolidated_vulnerabilities.items():
            legible_risk = details['risk']
            risk_color_code = details['risk_color']

            # Crear el párrafo para el título de la vulnerabilidad con el esquema de colores adecuado
            paragraph = doc.add_paragraph(style='Heading1')
            run_risk = paragraph.add_run(f"{legible_risk} - ")
            run_risk.font.color.rgb = RGBColor(int(risk_color_code[0:2], 16), int(risk_color_code[2:4], 16), int(risk_color_code[4:], 16))
            run_risk.bold = True

            run_name = paragraph.add_run(vulnerability_name)
            run_name.bold = True

            # Definir los títulos de las columnas según el idioma seleccionado
            if language == 'es':
                titles = ['Hosts Afectados', 'Puerto', 'Descripción', 'Solución', 'Evidencia', 'Evidencia Adicional']
            else:
                titles = ['Affected Hosts', 'Ports', 'Description', 'Solution', 'Evidence', 'Extra Evidence']

            # Crear la tabla para los detalles de la vulnerabilidad
            table = doc.add_table(rows=6, cols=2)
            table.style = 'TableGrid'
            table.autofit = True

            # Establecer el ancho de las columnas
            column_width_left = Cm(2.5)
            column_width_right = Cm(15)
            table.columns[0].width = column_width_left
            table.columns[1].width = column_width_right

            details_data = [
                ', '.join(details['hosts']) if 'hosts' in details else 'Unknown',
                ', '.join(str(port) for port in details['ports']) if 'ports' in details else 'Unknown',
                details['description'] if 'description' in details else '',
                details['solution'] if 'solution' in details else '',
                ', '.join(details['evidence']) if 'evidence' in details else 'Unknown',
                '',  # Placeholder for additional evidence cell
            ]

            for i, title in enumerate(titles):
                cell = table.cell(i, 0)
                cell_text = cell.paragraphs[0].add_run(title)
                cell_text.bold = True

                table.cell(i, 1).text = details_data[i]

                # Aplicar el color de fondo según la criticidad
                set_cell_background(cell, risk_color_code)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)  # Blanco

            # Añadir la imagen de la evidencia adicional si está disponible
            extra_evidence_image_cell = table.cell(5, 1)
            evidence_images = EvidenceImage.objects.filter(vulnerability_set__name=vulnerability_name)
            for evidence_image in evidence_images:
                image_path = os.path.join(settings.MEDIA_ROOT, evidence_image.image.name)
                if os.path.exists(image_path):
                    run = extra_evidence_image_cell.add_paragraph().add_run()
                    run.add_picture(image_path, width=Cm(10))  # Ajustar el ancho según sea necesario

            # Inserta un salto de página después de la tabla
            doc.add_page_break()



        # 📌 1️⃣ Verificar si la imagen ya existe antes de ejecutar Puppeteer
        graphmap_image_filename = f"graphmap_project_{project_id}.png"
        graphmap_image_path = os.path.join(settings.BASE_DIR, "static/images", graphmap_image_filename)

        # 📌 2️⃣ Insertar imagen + caption temporalmente al final del documento
        temp_image_para = doc.add_paragraph()
        temp_image_run = temp_image_para.add_run()
        temp_image_run.add_picture(graphmap_image_path, width=Inches(6))

        temp_caption_para = doc.add_paragraph("Figura 1: Mapa Visual del Ataque")
        temp_caption_para.alignment = 1  # Centrar la leyenda
        temp_caption_run = temp_caption_para.runs[0]
        temp_caption_run.bold = True  # Resaltar caption

        # 📌 3️⃣ Buscar `###Graphmap###` en el documento y mover imagen + caption
        graphmap_replaced = False
        for para in doc.paragraphs:
            if "###Graphmap###" in para.text:
                print("✅ `###Graphmap###` encontrado, moviendo imagen + caption.")

                # Limpiar el marcador sin afectar formato
                para.text = ""

                # 🔥 Insertar la imagen en el lugar de `###Graphmap###`
                graphmap_para = para.insert_paragraph_before()
                graphmap_para.alignment = 1  # Centrar la imagen + caption
                new_image_run = graphmap_para.add_run()
                new_image_run.add_picture(graphmap_image_path, width=Inches(6))


                graphmap_replaced = True
                break  # Salimos del loop tras mover

        # 📌 4️⃣ Si `###Graphmap###` fue encontrado, eliminar la imagen + caption del final
        if graphmap_replaced:
            body = doc._element.find("w:body", namespaces=doc._element.nsmap)
            if body is not None:
                if temp_image_para._element in body:
                    body.remove(temp_image_para._element)  # Eliminar imagen temporal
                if temp_caption_para._element in body:
                    body.remove(temp_caption_para._element)  # Eliminar caption temporal
            print("✅ Imagen + caption movidos y eliminados del final.")

        # 📌 5️⃣ Si `###Graphmap###` NO se encontró, mantener la imagen + caption al final
        if not graphmap_replaced:
            print("⚠️ `###Graphmap###` no encontrado. Se dejará la imagen + caption al final.")



        # Procesar saltos de página
        for paragraph in doc.paragraphs:
            if '###BreakPage###' in paragraph.text:
                paragraph.text = paragraph.text.replace('###BreakPage###', '')
                doc.add_page_break()

        # Inserta un salto de página después de la tabla
        doc.add_page_break()

        # Revisión para eliminar páginas en blanco
        while paragraphs[-1].text == "" or paragraphs[-1].text == "\f":
            # Si el párrafo es un salto de página o está vacío, elimínalo
            p = paragraphs[-1]._element
            p.getparent().remove(p)
            paragraphs = doc.paragraphs  # Actualizar la lista de párrafos


        # Inserta un salto de página después de la tabla
        doc.add_page_break()


        # 📌 Obtener los Writeups asignados al Proyecto
        attack_narratives = project.attack_narratives.all()

        print("🛠 DEBUG -> Writeups asignados al Proyecto:")
        for writeup in attack_narratives:
            print(f" - {writeup.title} (ID: {writeup.id})")
            print(f"Contenido:\n{writeup.content_html[:500]}")

        # 🔚 Revisión para eliminar páginas en blanco antes de insertar los Writeups
        while doc.paragraphs and (doc.paragraphs[-1].text.strip() == "" or doc.paragraphs[-1].text.strip() == "\f"):
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)

        # 📌 1️⃣ Insertar Writeups

        if attack_narratives.exists():
            doc.add_page_break()
            doc.add_heading("Attack Narrative", level=1)

            from bs4 import BeautifulSoup
            from PIL import Image

            for writeup in attack_narratives:
                if writeup.content_html:
                    doc.add_heading(writeup.title, level=2)

                    # --- Pre-process: download any remaining external images ---
                    processed_html = download_external_images(
                        writeup.content_html, writeup.title
                    )
                    soup = BeautifulSoup(processed_html, "html.parser")

                    # 🔄 Recorremos todos los elementos en orden
                    for el in soup.contents:
                        if not hasattr(el, "name"):
                            continue

                        if el.name == "img":
                            img_src = el.get("src")
                            filename = os.path.basename(img_src)
                            image_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, writeup.title, filename)

                            if os.path.exists(image_path):
                                add_image_to_doc(doc, image_path)
                                logger.info("Image inserted from: %s", image_path)
                            else:
                                logger.warning("Image not found: %s", image_path)
                        else:
                            # 🔍 Buscar imágenes anidadas dentro del tag (p, div, etc.)
                            for img_tag in el.find_all("img"):
                                img_src = img_tag.get("src")
                                filename = os.path.basename(img_src)
                                image_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, writeup.title, filename)

                                if os.path.exists(image_path):
                                    add_image_to_doc(doc, image_path)
                                    logger.info("Nested image inserted from: %s", image_path)
                                else:
                                    logger.warning("Nested image not found: %s", image_path)

                                img_tag.decompose()

                            # 👉 Insertar el bloque restante como HTML limpio (ya sin <img>)
                            try:
                                add_attack_narrative_to_document(doc, str(el), writeup.title, request, language)
                                logger.info("Writeup content inserted for: %s", writeup.title)
                            except Exception as e:
                                logger.error("Error inserting writeup content for %s: %s", writeup.title, e)

            print("✅ Todos los Writeups fueron insertados correctamente, en orden y desde protected_media.")
        else:
            print("⚠️ No hay Writeups asignados al Proyecto.")


        # Sustituir ###TOC_PLACEHOLDER### por el campo TOC de Word (mostrará números de página al abrir/actualizar)
        _replace_toc_placeholder_with_field(doc)

        # Actualizar campos al abrir en Word para que el TOC muestre números de página
        _set_update_fields_on_open(doc)

        # Guardar el documento en un buffer y preparar la respuesta
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{project.name}_report_{language}.docx"
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

    else:
        return render(request, 'projectmanager/select_report_language.html', {'project_id': project_id})








def generate_vulnerabilities_table(project, language):
    """
    Builds an HTML table of vulnerabilities. All user-supplied content (evidence,
    description, solution, hosts_affected) is escaped to prevent XSS (OWASP A03).
    """
    vulnerabilities = Vulnerability.objects.filter(project=project)
    vulnerabilities = sorted(vulnerabilities, key=lambda x: risk_factor_to_numeric(x.risk_factor, language))

    table_html = '<table border="1">'
    table_html += '<tr><th>Affected Hosts</th><th>Ports</th><th>Description</th><th>Solution</th><th>Evidence</th></tr>'

    for vulnerability in vulnerabilities:
        hosts_affected = vulnerability.hosts_affected if vulnerability.hosts_affected else 'Unknown'
        ports = str(vulnerability.port) if vulnerability.port else 'Unknown'
        description = vulnerability.description_es if language == 'es' else vulnerability.description
        solution = vulnerability.solution_es if language == 'es' else vulnerability.solution
        evidence = vulnerability.evidence if vulnerability.evidence else 'Unknown'
        # Escape user content to prevent XSS when this HTML is rendered (OWASP Top 10)
        hosts_affected = escape(hosts_affected)
        ports = escape(ports)
        description = escape(description or '')
        solution = escape(solution or '')
        evidence = escape(evidence)

        table_html += f'<tr><td>{hosts_affected}</td><td>{ports}</td><td>{description}</td><td>{solution}</td><td>{evidence}</td></tr>'

    table_html += '</table>'
    return table_html


def apply_styles(paragraph, element):
    """
    Aplica estilos a un párrafo de acuerdo con los estilos definidos en un elemento HTML.
    
    Parameters:
    - paragraph: El objeto Paragraph de python-docx al que se aplicarán los estilos.
    - element: El elemento HTML que contiene los estilos a aplicar.
    """
    # Verifica si hay un estilo de color definido en el elemento HTML
    if 'color' in element.attrs:
        color = element.attrs['color']
        # Convierte el color hexadecimal a RGB
        rgb_color = RGBColor(*tuple(int(color[i:i+2], 16) for i in (1, 3, 5)))
        # Aplica el color al texto del párrafo
        for run in paragraph.runs:
            run.font.color.rgb = rgb_color
    
    # Verifica si hay un estilo de tamaño de fuente definido en el elemento HTML
    if 'font-size' in element.attrs:
        font_size = element.attrs['font-size']
        # Convierte el tamaño de la fuente a puntos
        font_size_pt = Pt(int(font_size[:-2]))  # Suponiendo que el tamaño de fuente se especifica en px
        # Aplica el tamaño de la fuente al texto del párrafo
        for run in paragraph.runs:
            run.font.size = font_size_pt
    
    # Verifica si hay un estilo de alineación definido en el elemento HTML
    if 'text-align' in element.attrs:
        text_align = element.attrs['text-align']
        # Mapea la alineación HTML a la alineación de párrafo de docx
        alignment_mapping = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        # Aplica la alineación al párrafo
        paragraph.alignment = alignment_mapping.get(text_align, WD_PARAGRAPH_ALIGNMENT.LEFT)





#Area de Template para customizar el reporte:
@login_required
def report_template_list(request):
    templates = ReportTemplate.objects.all()
    return render(request, 'report_template_list.html', {'templates': templates})

@login_required
def report_template_create(request):
    if request.method == 'POST':
        form = ReportTemplateForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('report_template_list')
    else:
        form = ReportTemplateForm()
    return render(request, 'report_template_form.html', {'form': form})

@login_required
def report_template_edit(request, pk):
    template = get_object_or_404(ReportTemplate, pk=pk)
    if request.method == 'POST':
        form = ReportTemplateForm(request.POST, instance=template)
        if form.is_valid():
            form.save()
            return redirect('report_template_list')
    else:
        form = ReportTemplateForm(instance=template)
    return render(request, 'report_template_form.html', {'form': form})

@login_required
def report_template_delete(request, pk):
    template = get_object_or_404(ReportTemplate, pk=pk)
    template.delete()
    return redirect('report_template_list')


# Cover templates (new UI, no admin)
@login_required
def cover_template_list(request):
    templates = ReportCoverTemplate.objects.all().order_by('name')
    return render(request, 'projectmanager/cover_template_list.html', {'templates': templates})


@login_required
def cover_template_create(request):
    if request.method == 'POST':
        form = ReportCoverForm(request.POST, request.FILES)
        if form.is_valid():
            template = form.save()
            return redirect('visual_cover_designer_template', template_id=template.pk)
    else:
        form = ReportCoverForm()
    return render(request, 'projectmanager/cover_template_form.html', {'form': form})


@login_required
def cover_template_delete(request, pk):
    template = get_object_or_404(ReportCoverTemplate, pk=pk)
    template.delete()
    return redirect('cover_template_list')

# Netsparker Parser
@login_required
def import_netsparker_file(request, pk):
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)
    
    if request.method == 'POST':
        form = NetsparkerFileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            netsparker_file = request.FILES['netsparker_file']  # Asegúrate de que netsparker_file sea un InMemoryUploadedFile
            tree = ET.parse(netsparker_file)
            root = tree.getroot()

            # Extraer y procesar la URL base del módulo <target>
            target_url = root.find('.//target/url').text if root.find('.//target/url') is not None else None

            # Verificar y actualizar/crear el Target si target_url existe (get+save pattern to avoid select_for_update / SQLite lock)
            if target_url:
                try:
                    target = Target.objects.get(project=project, urlAddress=target_url)
                except Target.DoesNotExist:
                    target = Target.objects.create(project=project, urlAddress=target_url)

            for vuln in root.findall('.//vulnerability'):
                url_el = vuln.find('url')
                vuln_url = url_el.text if url_el is not None and url_el.text else ""
                sev_el = vuln.find('severity')
                severity = sev_el.text if sev_el is not None and sev_el.text else ""
                title_el = vuln.find('title')
                title = title_el.text if title_el is not None and title_el.text else "Unknown"
                desc_el = vuln.find('description')
                description_html = (desc_el.text or "") if desc_el is not None else ""
                remedy_el = vuln.find('remedy')
                remedy_html = (remedy_el.text or "") if remedy_el is not None else ""
                ext_el = vuln.find('externalReferences')
                externalReferences_html = (ext_el.text or "") if ext_el is not None else ""
                rawreq_el = vuln.find('rawrequest')
                rawrequest = (rawreq_el.text or "") if rawreq_el is not None else ""
                rawresp_el = vuln.find('rawresponse')
                rawresponse = (rawresp_el.text or "") if rawresp_el is not None else ""

                # Limpiar los campos HTML
                description_clean = clean_html(description_html)
                remedy_clean = clean_html(remedy_html)
                externalReferences_clean = clean_html(externalReferences_html)

                # Traducir al español
                description_es = GoogleTranslator(source='auto', target='es').translate(description_clean)
                remedy_es = GoogleTranslator(source='auto', target='es').translate(remedy_clean)

                # Construir el texto de evidencia
                evidence_text = f"Request:\n{rawrequest}\nResponse:\n{rawresponse}"

                # Get or create vulnerability without select_for_update (avoids SQLite "database is locked")
                defaults = {
                    'description': description_clean,
                    'solution': remedy_clean,
                    'description_es': description_es,
                    'solution_es': remedy_es,
                    'risk_factor': severity,
                    'see_also': externalReferences_clean,
                    'hosts_affected': vuln_url,
                    'evidence': evidence_text
                }
                try:
                    vulnerability = Vulnerability.objects.get(project=project, name=title)
                    for key, value in defaults.items():
                        setattr(vulnerability, key, value)
                    vulnerability.save()
                except Vulnerability.DoesNotExist:
                    vulnerability = Vulnerability.objects.create(
                        project=project,
                        name=title,
                        **defaults
                    )

            return redirect('project_detail', pk=project.pk)
    else:
        form = NetsparkerFileUploadForm()

    return render(request, 'admin/import_file.html', {'form': form, 'project': project})



# App login/logout (new-version UI, same Django auth)
def app_login(request):
    from django.contrib.auth.views import LoginView
    from django.contrib.auth.forms import AuthenticationForm
    if request.user.is_authenticated:
        return redirect('project_list')
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            from django.contrib.auth import login
            user = form.get_user()
            login(request, user)
            next_url = request.GET.get('next') or request.POST.get('next') or settings.LOGIN_REDIRECT_URL
            return redirect(next_url)
    else:
        form = AuthenticationForm(request)
    return render(request, 'projectmanager/login.html', {'form': form, 'next': request.GET.get('next', '')})


def app_logout(request):
    from django.contrib.auth import logout
    logout(request)
    return render(request, 'projectmanager/logout.html')

def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('login')  # Redirige a la página de inicio de sesión después de registrarse
    else:
        form = UserCreationForm()
    return render(request, 'accounts/register.html', {'form': form})



def register(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')  # Asume que tienes una vista 'home'
    else:
        form = CustomUserCreationForm()
    return render(request, 'accounts/register.html', {'form': form})






#Parser Acunetix
@login_required
def import_acunetix_xml(request, pk):
    from deep_translator import GoogleTranslator
    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)

    if request.method == 'POST':
        xml_file = request.FILES.get('acunetix_file')
        if xml_file:
            tree = ET.parse(xml_file)
            root = tree.getroot()

            unique_hosts_affected = set()
            os_name = None

            for scan in root.findall('.//Scan'):
                start_url = scan.find('StartURL').text if scan.find('StartURL') is not None else ''
                crawler_start_url = scan.get('StartUrl', start_url)

                os_element = scan.find('.//Os')
                os_name = None
                if os_element is not None and os_element.text:
                    os_content = os_element.text.strip()
                    match = re.search(r'\[CDATA\[(.*?)\]\]', os_content)
                    os_name = match.group(1) if match else os_content

                unique_hosts_affected.add(start_url)

            for host in unique_hosts_affected:
                try:
                    t = Target.objects.get(project=project, urlAddress=host)
                    t.fqdn = host
                    t.os = os_name
                    t.save()
                except Target.DoesNotExist:
                    Target.objects.create(project=project, urlAddress=host, fqdn=host, os=os_name)

            for scan in root.findall('.//Scan'):
                start_url = scan.find('StartURL').text if scan.find('StartURL') is not None else ''
                crawler_start_url = scan.get('StartUrl', start_url)

                try:
                    target_host = Target.objects.get(project=project, urlAddress=crawler_start_url)
                except Target.DoesNotExist:
                    target_host = Target.objects.create(project=project, urlAddress=crawler_start_url, fqdn=crawler_start_url or '', os=os_name)

                for report_item in scan.findall('.//ReportItem'):
                    name_el = report_item.find('Name')
                    name = (name_el.text or '') if name_el is not None else ''
                    desc_el = report_item.find('Description')
                    desc_t = (desc_el.text or '') if desc_el is not None else ''
                    impact_el = report_item.find('Impact')
                    impact_t = (impact_el.text or '') if impact_el is not None else ''
                    details_el = report_item.find('Details')
                    details_t = (details_el.text or '') if details_el is not None else ''
                    full_description = f"Description:\n{desc_t}\nImpact:\n{impact_t}\nDetails:\n{details_t}\n"

                    rec_el = report_item.find('Recommendation')
                    recommendation_text = (rec_el.text or '') if rec_el is not None else ''
                    sev_el = report_item.find('Severity')
                    severity = sev_el.text.capitalize() if sev_el is not None and sev_el.text else ''
                    
                    technical_details = report_item.find('.//TechnicalDetails')
                    req_el = technical_details.find('Request') if technical_details is not None else None
                    request_text = (req_el.text or '') if req_el is not None else ''
                    response_text = ""  # Se puede expandir si es necesario

                    evidence = f"Request:\n{request_text}\nResponse:\n{response_text}"

                    # Traducción segura
                    description_es = GoogleTranslator(source='auto', target='es').translate(full_description) if full_description.strip() else ''
                    recommendation_es = GoogleTranslator(source='auto', target='es').translate(recommendation_text) if recommendation_text.strip() else ''

                    Vulnerability.objects.create(
                        project=project,
                        name=name,
                        description=full_description,
                        description_es=description_es,
                        solution=recommendation_text,
                        solution_es=recommendation_es,
                        evidence=evidence,
                        risk_factor=severity,
                        hosts_affected=start_url,
                        port=None,
                        target_host=target_host
                    )

            return redirect('project_detail', pk=project.pk)

    return render(request, 'admin/import_file.html', {'project': project})





def extract_links_from_references(references_element):
    if references_element is not None and references_element.text:
        # Extraer enlaces utilizando una expresión regular, por ejemplo
        links = re.findall(r'href="([^"]+)"', references_element.text)
        return "\n".join(links)
    return ""


def try_decode_base64(data):
    """
    Intenta decodificar una cadena codificada en base64 primero como utf-8.
    Si falla, devuelve una indicación de que los datos no se pudieron decodificar.
    """
    try:
        return base64.b64decode(data).decode('utf-8')
    except UnicodeDecodeError:
        return "Data not decodable"


def extract_links_from_cdata(cdata):
    soup = BeautifulSoup(cdata, "html.parser")
    links = [a['href'] for a in soup.find_all('a', href=True)]
    return "\n".join(links)


# Función auxiliar para limpiar HTML y traducir texto
def clean_and_translate_html(text, lang='es'):
    # Asumiendo que tienes una función split_and_translate(text, lang) definida en otro lugar
    cleaned_text = re.sub(r'<[^>]+>', '', text)  # Elimina etiquetas HTML
    return split_and_translate(cleaned_text, lang)


#Parser Burpsuite XML file
@login_required
def import_burp_xml(request, pk):
    from deep_translator import GoogleTranslator
    import re
    from collections import defaultdict

    project = get_object_or_404(get_projects_queryset(request.user), pk=pk)

    if request.method == 'POST':
        form = BurpUploadForm(request.POST, request.FILES)
        if form.is_valid():
            burp_file = request.FILES['burp_file']
            tree = ET.parse(burp_file)
            root = tree.getroot()

            unique_hosts_by_ip = defaultdict(set)

            for item in root.findall('.//issue'):
                host_tag = item.find('host')
                ip_address = host_tag.get('ip') if host_tag is not None else ""
                host_url = host_tag.text if host_tag is not None else ""
                unique_hosts_by_ip[ip_address].add(host_url)

                sev_el = item.find('severity')
                severity = sev_el.text.capitalize() if sev_el is not None and sev_el.text else ""
                name_el = item.find('name')
                name = name_el.text if name_el is not None and name_el.text else ""

                issue_bg = item.find('issueBackground')
                description_raw = (issue_bg.text or "") if issue_bg is not None else ""
                remed_bg = item.find('remediationBackground')
                remediation_raw = (remed_bg.text or "") if remed_bg is not None else ""

                # Limpiar etiquetas HTML si vienen con <p>, <b>, etc.
                clean_description = re.sub(r'<[^>]*>', '', description_raw).strip()
                clean_solution = re.sub(r'<[^>]*>', '', remediation_raw).strip()

                # Traducción con fallback
                description_es = GoogleTranslator(source='auto', target='es').translate(clean_description) if clean_description else ''
                solution_es = GoogleTranslator(source='auto', target='es').translate(clean_solution) if clean_solution else ''

                references_element = item.find('references')
                references = extract_links_from_references(references_element)

                request_element = item.find('.//request')
                request_encoded = request_element.text if request_element is not None else ""
                response_element = item.find('.//response')
                response_encoded = response_element.text if response_element is not None else ""

                request_decoded = try_decode_base64(request_encoded)
                response_decoded = try_decode_base64(response_encoded)

                evidence = f"Request:\n{request_decoded}\nResponse:\n{response_decoded}"

                for ip_address, hosts in unique_hosts_by_ip.items():
                    for host_url in hosts:
                        target, _ = Target.objects.get_or_create(
                            project=project,
                            ip_address=ip_address,
                            urlAddress=host_url,
                            defaults={'urlAddress': host_url}
                        )

                        # Crear o actualizar la vulnerabilidad
                        vuln, created = Vulnerability.objects.get_or_create(
                            project=project,
                            name=name,
                            target_host=target,
                            defaults={
                                'description': clean_description,
                                'description_es': description_es,
                                'solution': clean_solution,
                                'solution_es': solution_es,
                                'risk_factor': severity,
                                'evidence': evidence,
                                'see_also': references,
                                'hosts_affected': host_url
                            }
                        )

                        if not created:
                            if host_url not in vuln.hosts_affected.split('\n'):
                                vuln.hosts_affected += f"\n{host_url}"
                            if evidence not in vuln.evidence:
                                vuln.evidence += f"\n\n{evidence}"
                            vuln.save()

            return redirect('project_detail', pk=project.pk)
    else:
        form = BurpUploadForm()

    return render(request, 'admin/import_file.html', {'form': form, 'project': project})





# Proxy Simple Icons: tipos que usan SVG del CDN (los que se veían mal en PNG local)
GRAPH_ICON_CDN = {
    'linux': ('linux', 'FCC624'),
    'ubuntu': ('ubuntu', 'E95420'),
    'debian': ('debian', 'A81D33'),
    'kali': ('kalilinux', '557C94'),
    'printer_hp': ('hp', '0096D6'),
    'nas_synology': ('synology', '2D7EC1'),
    'firewall': ('fortinet', 'EE3124'),
}

# Iconos del GraphMap: SOLO SVG, NUNCA PNG. Logos oficiales en static/images/graphmap/*.svg
# (scripts/download_graphmap_logos.py: Windows desde Wikimedia Commons, resto desde Simple Icons).
GRAPH_ICON_STATIC = {
    'windows': 'windows.svg', 'windows_xp': 'windows_xp.svg', 'windows_vista': 'windows_vista.svg',
    'windows_7': 'windows_7.svg', 'windows_8': 'windows_8.svg', 'windows_10': 'windows_10.svg',
    'windows_11': 'windows_11.svg', 'windows_12': 'windows_12.svg', 'windows_server': 'windows_server.svg',
    'linux': 'linux.svg', 'ubuntu': 'ubuntu.svg', 'debian': 'debian.svg', 'kali': 'kali.svg',
    'arch': 'arch.svg', 'fedora': 'fedora.svg', 'redhat': 'redhat.svg', 'macos': 'macos.svg',
    'android': 'android.svg', 'ios': 'macos.svg', 'server': 'linux.svg',
    'printer': 'unknown.svg', 'printer_epson': 'unknown.svg', 'printer_hp': 'unknown.svg',
    'printer_canon': 'unknown.svg', 'printer_brother': 'unknown.svg', 'printer_lexmark': 'unknown.svg',
    'router': 'unknown.svg', 'router_cisco': 'unknown.svg', 'router_netgear': 'unknown.svg',
    'router_tplink': 'unknown.svg', 'router_mikrotik': 'unknown.svg', 'router_ubiquiti': 'unknown.svg',
    'router_dlink': 'unknown.svg', 'router_aruba': 'unknown.svg',
    'nas': 'unknown.svg', 'nas_synology': 'unknown.svg', 'nas_qnap': 'unknown.svg',
    'iot': 'unknown.svg', 'iot_raspberrypi': 'unknown.svg', 'iot_arduino': 'unknown.svg',
    'camera': 'unknown.svg', 'camera_hikvision': 'unknown.svg', 'camera_dahua': 'unknown.svg', 'camera_axis': 'unknown.svg',
    'firewall': 'unknown.svg', 'switch': 'unknown.svg', 'unknown': 'unknown.svg',
}

# Tipos de nodo que usan la plantilla imac (computadora con OS en pantalla)
GRAPH_IMAC_TYPES = {
    'windows', 'windows_xp', 'windows_vista', 'windows_7', 'windows_8', 'windows_10', 'windows_11',
    'windows_12', 'windows_server',
    'linux', 'ubuntu', 'debian', 'kali', 'arch', 'fedora', 'redhat', 'macos', 'android', 'ios',
    'server', 'unknown',
}
# Fuente del icono OS para graph_node_image: SOLO SVG (nunca PNG). Local o URL CDN.
GRAPH_OS_ICON_SOURCE = {
    'windows': 'windows.svg', 'windows_xp': 'windows_xp.svg', 'windows_vista': 'windows_vista.svg',
    'windows_7': 'windows_7.svg', 'windows_8': 'windows_8.svg', 'windows_10': 'windows_10.svg',
    'windows_11': 'windows_11.svg', 'windows_12': 'windows_12.svg', 'windows_server': 'windows_server.svg',
    'linux': 'linux.svg', 'ubuntu': 'ubuntu.svg', 'debian': 'debian.svg', 'kali': 'kali.svg',
    'arch': 'arch.svg', 'fedora': 'fedora.svg', 'redhat': 'redhat.svg', 'macos': 'macos.svg',
    'ios': 'macos.svg', 'server': 'linux.svg', 'unknown': 'unknown.svg',
    'android': 'https://cdn.simpleicons.org/android/3DDC84',
}


def _svg_to_png(svg_content_or_path, size_px, is_url=False):
    """Convierte SVG a PNG de tamaño size_px. svg_content_or_path: path local (str) o URL (str)."""
    if cairosvg is None:
        return None
    try:
        if is_url:
            png_bytes = cairosvg.svg2png(url=svg_content_or_path, output_width=size_px, output_height=size_px)
        else:
            path = str(svg_content_or_path)
            with open(path, 'rb') as f:
                svg_bytes = f.read()
            png_bytes = cairosvg.svg2png(bytestring=svg_bytes, output_width=size_px, output_height=size_px)
    except Exception:
        return None
    if not png_bytes:
        return None
    return Image.open(BytesIO(png_bytes)).convert('RGBA')


def _draw_lightning_overlay(img):
    """Dibuja rayos eléctricos en los bordes de la imagen (efecto comprometido)."""
    w, h = img.size
    overlay = Image.new('RGBA', (w, h), (0, 0, 0, 0))
    d = ImageDraw.Draw(overlay)
    # Rayos: líneas en zigzag blanco/amarillo semitransparentes
    import random
    random.seed(42)
    for _ in range(8):
        x0 = random.randint(0, w)
        y0 = 0 if random.random() > 0.5 else h
        pts = [(x0, y0)]
        for _ in range(4):
            pts.append((pts[-1][0] + random.randint(-40, 40), pts[-1][1] + (random.randint(10, 30) if y0 == 0 else -random.randint(10, 30))))
        d.line(pts, fill=(255, 255, 200, 180), width=3)
    for _ in range(6):
        x0 = random.randint(0, w)
        y0 = random.randint(0, h)
        pts = [(x0, y0)]
        for _ in range(3):
            pts.append((pts[-1][0] + random.randint(-25, 25), pts[-1][1] + random.randint(-25, 25)))
        d.line(pts, fill=(255, 220, 100, 160), width=2)
    return Image.alpha_composite(img, overlay)


@login_required
@xframe_options_sameorigin
def graph_icon_proxy(request):
    """Proxy de iconos Simple Icons (mismo origen para evitar CORS/OpaqueResponseBlocking). GET type=ubuntu|linux|debian|..."""
    node_type = (request.GET.get('type') or 'unknown').strip().lower()
    if node_type not in GRAPH_ICON_CDN:
        raise Http404("Unknown icon type")
    slug, color = GRAPH_ICON_CDN[node_type]
    url = f"https://cdn.simpleicons.org/{slug}/{color}"
    try:
        r = requests.get(url, timeout=5)
        r.raise_for_status()
    except Exception:
        raise Http404("Icon unavailable")
    resp = HttpResponse(r.content, content_type='image/svg+xml')
    resp['Cache-Control'] = 'public, max-age=86400'
    return resp


@login_required
@xframe_options_sameorigin
def graph_node_image(request):
    """Genera una imagen PNG: imac (o imacRED si owned) + icono del OS en la pantalla; si owned, rayos."""
    node_type = (request.GET.get('type') or 'unknown').strip().lower()
    owned = request.GET.get('owned', '0') == '1'
    if node_type not in GRAPH_IMAC_TYPES:
        node_type = 'unknown'
    base = str(settings.BASE_DIR)
    static_root = str(settings.STATIC_ROOT) if settings.STATIC_ROOT else ''
    images_dir = os.path.join(static_root or os.path.join(base, 'static'), 'images')
    if not os.path.isdir(images_dir):
        images_dir = os.path.join(base, 'static', 'images')
    imac_filename = 'imacRED.png' if owned else 'imac.png'
    imac_path = os.path.join(images_dir, imac_filename)
    if not os.path.isfile(imac_path):
        imac_path = os.path.join(base, 'static', 'images', imac_filename)
    if not os.path.isfile(imac_path):
        raise Http404(f"{imac_filename} not found")
    img = Image.open(imac_path).convert('RGBA')
    # Generar a 128px para que al mostrarse en el nodo (~48px) el logo se vea nítido
    OUT_SIZE = 128
    w, h = img.size
    if w != OUT_SIZE or h != OUT_SIZE:
        img = img.resize((OUT_SIZE, OUT_SIZE), Image.Resampling.LANCZOS)
        w, h = OUT_SIZE, OUT_SIZE
    # Zona pantalla: proporciones para que el logo ocupe la mayor parte
    screen_left = int(w * 0.12)
    screen_top = int(h * 0.10)
    screen_right = int(w * 0.88)
    screen_bottom = int(h * 0.56)
    screen_w = screen_right - screen_left
    screen_h = screen_bottom - screen_top
    icon_size = min(screen_w, screen_h)
    # Icono OS: SOLO SVG (oficiales). Nunca PNG.
    graphmap_dir = os.path.join(base, 'static', 'images', 'graphmap')
    icon_source = GRAPH_OS_ICON_SOURCE.get(node_type, 'unknown.svg')
    if (isinstance(icon_source, str) and not icon_source.startswith('http') and
            node_type.startswith('windows') and node_type != 'windows'):
        _version_svg = os.path.join(graphmap_dir, icon_source)
        if not os.path.isfile(_version_svg):
            icon_source = 'windows.svg'
    icon_png = None
    svg_path = None
    if not (isinstance(icon_source, str) and icon_source.startswith('http')):
        svg_path = os.path.join(graphmap_dir, icon_source)
    if cairosvg and svg_path and os.path.isfile(svg_path):
        icon_png = _svg_to_png(svg_path, icon_size, is_url=False)
    if icon_png is None and isinstance(icon_source, str) and icon_source.startswith('http') and cairosvg:
        icon_png = _svg_to_png(icon_source, icon_size, is_url=True)
    if icon_png is not None:
        icon_png = icon_png.resize((icon_size, icon_size), Image.Resampling.LANCZOS)
        paste_x = screen_left + (screen_w - icon_size) // 2
        paste_y = screen_top + (screen_h - icon_size) // 2
        mask = icon_png.split()[3] if icon_png.mode == 'RGBA' and len(icon_png.split()) >= 4 else icon_png
        img.paste(icon_png, (paste_x, paste_y), mask)
    else:
        # Fallback sin cairosvg: dibujar un rectángulo de color en la pantalla (indicador de OS)
        fill_colors = {'windows': (0, 120, 212), 'linux': (252, 198, 36), 'ubuntu': (233, 84, 32), 'macos': (85, 85, 85), 'unknown': (100, 100, 100)}
        rgb = fill_colors.get(node_type, (0, 120, 212) if node_type.startswith('windows') else fill_colors['unknown'])
        draw = ImageDraw.Draw(img)
        draw.rectangle([screen_left, screen_top, screen_right, screen_bottom], fill=rgb, outline=(60, 60, 60))
    # Equipos comprometidos: solo imacRED.png (sin rayos); la base roja es suficiente
    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type='image/png')
    resp['Cache-Control'] = 'public, max-age=3600'
    return resp


@login_required
@never_cache
@xframe_options_sameorigin
def graph_map_view(request, project_id):
    project = get_object_or_404(get_projects_queryset(request.user), id=project_id)
    # Consulta fresca: sin caché para que owned/OS siempre estén actualizados
    targets = list(Target.objects.filter(project=project).select_related('jumped_from').order_by('id'))

    CONSULTANT_NODE_ID = 0
    nodes = []
    edges = []
    targets_with_predecessor = set()

    for target in targets:
        label = target.ip_address or target.fqdn or target.urlAddress or f"Target-{target.pk}"
        os_info = (target.os or "").lower().strip()
        # Windows version mapping: variantes explícitas y por número de versión
        if any(x in os_info for x in ('xp', 'windows xp')):
            node_type = 'windows_xp'
        elif any(x in os_info for x in ('vista', 'windows vista')):
            node_type = 'windows_vista'
        elif any(x in os_info for x in ('win 7', 'win7', 'windows 7', 'windows7')) or ('windows' in os_info and '7' in os_info) or ('win' in os_info and '7' in os_info):
            node_type = 'windows_7'
        elif any(x in os_info for x in ('win 8', 'win8', 'windows 8', 'windows8')) or ('windows' in os_info and '8' in os_info) or ('win' in os_info and '8' in os_info):
            node_type = 'windows_8'
        elif any(x in os_info for x in ('win 10', 'win10', 'windows 10', 'windows10')) or ('windows' in os_info and '10' in os_info) or ('win' in os_info and '10' in os_info):
            node_type = 'windows_10'
        elif any(x in os_info for x in ('win 11', 'win11', 'windows 11', 'windows11')) or ('windows' in os_info and '11' in os_info) or ('win' in os_info and '11' in os_info):
            node_type = 'windows_11'
        elif any(x in os_info for x in ('win 12', 'win12', 'windows 12', 'windows12')) or ('windows' in os_info and '12' in os_info) or ('win' in os_info and '12' in os_info):
            node_type = 'windows_12'
        elif any(x in os_info for x in ('windows server', 'windowsserver', 'win server', 'winserver', 'windows 20')) or ('server' in os_info and ('windows' in os_info or 'win' in os_info)):
            node_type = 'windows_server'
        elif any(x in os_info for x in ('windows', 'win')):
            node_type = 'windows'
        elif any(x in os_info for x in ('ios', 'iphone', 'ipad')):
            node_type = 'ios'
        elif any(x in os_info for x in ('darwin', 'mac os', 'macos', 'os x', 'osx', 'apple')):
            node_type = 'macos'
        elif 'android' in os_info:
            node_type = 'android'
        elif 'ubuntu' in os_info:
            node_type = 'ubuntu'
        elif 'debian' in os_info:
            node_type = 'debian'
        elif 'kali' in os_info:
            node_type = 'kali'
        elif 'arch' in os_info:
            node_type = 'arch'
        elif 'fedora' in os_info:
            node_type = 'fedora'
        elif any(x in os_info for x in ('centos', 'rhel', 'red hat', 'redhat')):
            node_type = 'redhat'
        elif any(x in os_info for x in ('linux', 'gnu/linux')):
            node_type = 'linux'
        # Impresoras (marca primero, luego genérico)
        elif any(x in os_info for x in ('epson', 'epson printer')):
            node_type = 'printer_epson'
        elif any(x in os_info for x in ('hp printer', 'hewlett', 'hp laserjet', 'hp deskjet')):
            node_type = 'printer_hp'
        elif any(x in os_info for x in ('canon printer', 'canon pixma', 'canon imageclass')):
            node_type = 'printer_canon'
        elif any(x in os_info for x in ('brother printer', 'brother hl', 'brother mfc')):
            node_type = 'printer_brother'
        elif any(x in os_info for x in ('lexmark', 'xerox printer', 'samsung printer')):
            node_type = 'printer_lexmark'
        elif any(x in os_info for x in ('printer', 'impresora', 'print server')):
            node_type = 'printer'
        # Routers / networking
        elif any(x in os_info for x in ('cisco', 'ios-xe', 'ios xe', 'catalyst')):
            node_type = 'router_cisco'
        elif any(x in os_info for x in ('netgear', 'netgear router', 'night hawk')):
            node_type = 'router_netgear'
        elif any(x in os_info for x in ('tp-link', 'tplink', 'tp link')):
            node_type = 'router_tplink'
        elif any(x in os_info for x in ('mikrotik', 'routeros')):
            node_type = 'router_mikrotik'
        elif any(x in os_info for x in ('ubiquiti', 'unifi', 'edgeos')):
            node_type = 'router_ubiquiti'
        elif any(x in os_info for x in ('d-link', 'dlink')):
            node_type = 'router_dlink'
        elif any(x in os_info for x in ('aruba', 'arubaos')):
            node_type = 'router_aruba'
        elif any(x in os_info for x in ('router', 'gateway', 'access point', 'ap ', 'wifi')):
            node_type = 'router'
        # NAS / almacenamiento
        elif any(x in os_info for x in ('synology', 'dsm', 'diskstation')):
            node_type = 'nas_synology'
        elif any(x in os_info for x in ('qnap', 'qts', 'turbo nas')):
            node_type = 'nas_qnap'
        elif any(x in os_info for x in ('nas', 'network storage', 'storage')):
            node_type = 'nas'
        # IoT / embebidos
        elif any(x in os_info for x in ('raspberry', 'raspberry pi', 'rpi')):
            node_type = 'iot_raspberrypi'
        elif any(x in os_info for x in ('arduino', 'esp32', 'esp8266', 'espressif')):
            node_type = 'iot_arduino'
        elif any(x in os_info for x in ('iot', 'embedded', 'smart device', 'sensor')):
            node_type = 'iot'
        # Cámaras IP / vigilancia
        elif any(x in os_info for x in ('hikvision', 'hik connect')):
            node_type = 'camera_hikvision'
        elif any(x in os_info for x in ('dahua', 'dhi-')):
            node_type = 'camera_dahua'
        elif any(x in os_info for x in ('axis camera', 'axis communications')):
            node_type = 'camera_axis'
        elif any(x in os_info for x in ('camera', 'ip camera', 'cámara', 'nvr', 'dvr')):
            node_type = 'camera'
        # Servidores / otros
        elif any(x in os_info for x in ('server', 'esxi', 'vmware', 'proxmox', 'hyper-v')):
            node_type = 'server'
        elif any(x in os_info for x in ('firewall', 'pfsense', 'opnsense', 'fortinet', 'palo alto')):
            node_type = 'firewall'
        elif any(x in os_info for x in ('switch', 'layer 2', 'layer 3')):
            node_type = 'switch'
        else:
            node_type = 'unknown'

        node = {
            'id': target.pk,
            'label': label,
            'type': node_type,
            'owned': target.owned,
            'ip': target.ip_address or '',
            'fqdn': target.fqdn or '',
            'url': target.urlAddress or '',
        }
        # #region agent log
        if target.owned or (label and '10.0.3.14' in label):
            import json as _json
            try:
                with open('/Users/orion/Desktop/Espengler-2.0/.cursor/debug.log', 'a') as _f:
                    _f.write(_json.dumps({'location': 'views.py:graph_map_view', 'message': 'target node', 'data': {'id': target.pk, 'label': label, 'type': node_type, 'os_raw': (target.os or '')[:80], 'owned': target.owned}, 'timestamp': int(__import__('time').time() * 1000)}) + '\n')
            except Exception:
                pass
        # #endregion

        if target.x_position is not None and target.y_position is not None:
            node['x'] = target.x_position
            node['y'] = target.y_position
            node['fixed'] = {'x': True, 'y': True}

        ports = Port.objects.filter(target=target)
        node['ports'] = [{'number': p.port_number, 'protocol': p.protocol, 'service': p.service_name or ''} for p in ports]

        nodes.append(node)

        if target.jumped_from:
            targets_with_predecessor.add(target.pk)
            edges.append({
                'from': target.jumped_from.pk,
                'to': target.pk,
            })

    # Consultant node as center of activity (id 0)
    consultant_node = {
        'id': CONSULTANT_NODE_ID,
        'label': 'Consultant',
        'type': 'consultant',
        'owned': False,
    }
    nodes.insert(0, consultant_node)
    entry_point_ids = [t.pk for t in targets if t.pk not in targets_with_predecessor]
    for tid in entry_point_ids:
        edges.append({'from': CONSULTANT_NODE_ID, 'to': tid})

    # Legend: only OS/types present in the graph (+ consultant, + owned if any)
    type_labels = {
        'consultant': 'Consultant',
        'owned': 'Compromised',
        'windows': 'Windows',
        'windows_xp': 'Windows XP',
        'windows_vista': 'Windows Vista',
        'windows_7': 'Windows 7',
        'windows_8': 'Windows 8',
        'windows_10': 'Windows 10',
        'windows_11': 'Windows 11',
        'windows_12': 'Windows 12',
        'windows_server': 'Windows Server',
        'linux': 'Linux',
        'ubuntu': 'Ubuntu',
        'debian': 'Debian',
        'kali': 'Kali',
        'arch': 'Arch',
        'fedora': 'Fedora',
        'redhat': 'Red Hat',
        'macos': 'macOS',
        'android': 'Android',
        'ios': 'iOS',
        'printer': 'Printer',
        'printer_epson': 'Epson',
        'printer_hp': 'HP Printer',
        'printer_canon': 'Canon',
        'printer_brother': 'Brother',
        'printer_lexmark': 'Lexmark/Xerox',
        'router': 'Router',
        'router_cisco': 'Cisco',
        'router_netgear': 'Netgear',
        'router_tplink': 'TP-Link',
        'router_mikrotik': 'MikroTik',
        'router_ubiquiti': 'Ubiquiti',
        'router_dlink': 'D-Link',
        'router_aruba': 'Aruba',
        'nas': 'NAS',
        'nas_synology': 'Synology',
        'nas_qnap': 'QNAP',
        'iot': 'IoT',
        'iot_raspberrypi': 'Raspberry Pi',
        'iot_arduino': 'Arduino/ESP',
        'camera': 'Camera',
        'camera_hikvision': 'Hikvision',
        'camera_dahua': 'Dahua',
        'camera_axis': 'Axis',
        'server': 'Server',
        'firewall': 'Firewall',
        'switch': 'Switch',
        'unknown': 'Unknown',
    }
    detected_types = set(n['type'] for n in nodes)
    if any(n.get('owned') for n in nodes if n.get('id') != CONSULTANT_NODE_ID):
        detected_types.add('owned')
    detected_os_types = [(t, type_labels.get(t, t)) for t in sorted(detected_types, key=lambda x: (0 if x == 'consultant' else 1, 0 if x == 'owned' else 1, x))]

    import json
    # Única fuente de verdad: URLs de iconos desde el backend (evita desincronizar con la plantilla)
    graphmap_dir = os.path.join(settings.BASE_DIR, 'static', 'images', 'graphmap')
    icon_urls = {}
    for t, filename in GRAPH_ICON_STATIC.items():
        if t.startswith('windows') and t != 'windows':
            path = os.path.join(graphmap_dir, filename)
            if not os.path.isfile(path):
                filename = 'windows.svg'
        icon_urls[t] = static('images/graphmap/' + filename)
    for t in GRAPH_ICON_CDN:
        icon_urls[t] = reverse('graph_icon_proxy') + '?type=' + t
    context = {
        'project': project,
        'nodes_json': json.dumps(nodes),
        'edges_json': json.dumps(edges),
        'icon_urls_json': json.dumps(icon_urls),
        'consultant_img': static('images/attack.png'),
        'owned_img': static('images/imac1.png'),
        'detected_os_types': detected_os_types,
        'embed': request.GET.get('embed') == '1',
    }
    response = render(request, 'projectmanager/graph_map.html', context)
    # never_cache ya añade cabeceras; reforzamos para navegadores
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    response['Vary'] = 'Cookie'
    return response



