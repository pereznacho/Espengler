# ProjectManager/reports.py
from django.shortcuts import render, get_object_or_404, redirect
from .models import Project, Vulnerability
from docx import Document
from docx.shared import Inches
from django.http import HttpResponse
import io
from google_trans_new import google_translator as GoogleTranslator


def generate_report(request, project_id):
    if request.method == 'POST':
        language = request.POST.get('language', 'EN')
        project = get_object_or_404(Project, pk=project_id)
        vulnerabilities = Vulnerability.objects.filter(project=project)

        doc = Document()
        doc.add_heading(project.nombre, 0)

        for vulnerability in vulnerabilities:
            doc.add_heading(f'Vulnerabilidad: {vulnerability.name}', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'TableGrid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Field'
            hdr_cells[1].text = 'Detail'

            description_translation = ''
            if vulnerability.description:
                # Utiliza GoogleTranslator para traducir el texto al idioma deseado
                description_translation = GoogleTranslator(source='en', target=language).translate(vulnerability.description)

            solution_translation = ''
            if vulnerability.solution:
                # Utiliza GoogleTranslator para traducir el texto al idioma deseado
                solution_translation = GoogleTranslator(source='en', target=language).translate(vulnerability.solution)

            data = [
                ('Detail', vulnerability.name),
                ('Solution', solution_translation if language == 'ES' else vulnerability.solution),
                ('Hosts Affected', vulnerability.hosts_affected if vulnerability.hosts_affected else 'Unknown'),
                ('Description', description_translation if language == 'ES' else vulnerability.description),
            ]

            for label, value in data:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(value)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{project.nombre}_report.docx"
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    else:
        return render(request, 'projectmanager/select_report_language.html', {'project_id': project_id})
