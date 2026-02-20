"""
image_utils.py — Centralized image handling for Espengler report generation.

Responsibilities:
  A) Download external images from HTML, store in protected_media/<writeup_name>/, rewrite HTML.
  B) Insert images into DOCX respecting document margins (no overflow).
"""

import os
import io
import re
import base64
import hashlib
import logging
import tempfile
import mimetypes
from urllib.parse import urlparse

import requests
from PIL import Image
from bs4 import BeautifulSoup

from django.conf import settings
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("espengler.images")


# ---------------------------------------------------------------------------
# A) External image downloader + HTML rewriter
# ---------------------------------------------------------------------------

def download_external_images(html_content, writeup_name):
    """
    Scan HTML for <img> tags with external (http/https) src attributes.
    Download each image into protected_media/<writeup_name>/ (avoiding duplicates),
    then rewrite the src to reference the local file.

    Returns:
        str: Updated HTML with local image paths.
    """
    if not html_content:
        return html_content

    if not writeup_name or writeup_name.strip() == "":
        writeup_name = "Unknown_Writeup"

    # Sanitize folder name
    safe_folder = writeup_name.replace(" ", "_")
    protected_path = os.path.join(settings.PROTECTED_MEDIA_ROOT, safe_folder)
    os.makedirs(protected_path, exist_ok=True)

    soup = BeautifulSoup(html_content, "html.parser")
    images_processed = 0

    for img_tag in soup.find_all("img"):
        src = img_tag.get("src", "")

        # Only process external URLs
        if not (src.startswith("http://") or src.startswith("https://")):
            continue

        # Skip if it's already pointing to our own server
        if "/protected_media/" in src:
            continue

        try:
            local_path = _download_and_save(src, protected_path)
            if local_path:
                # Rewrite src to local relative path for DOCX generation
                filename = os.path.basename(local_path)
                img_tag["src"] = f"/protected_media/{safe_folder}/{filename}"
                images_processed += 1
                logger.info(
                    "Downloaded external image: %s -> %s (writeup=%s)",
                    src, local_path, writeup_name
                )
        except Exception as e:
            logger.error(
                "Failed to download external image %s for writeup '%s': %s",
                src, writeup_name, e
            )

    if images_processed > 0:
        logger.info(
            "Processed %d external images for writeup '%s'",
            images_processed, writeup_name
        )

    return str(soup)


def _download_and_save(url, dest_dir):
    """
    Download an image from `url` and save it to `dest_dir`.
    Uses a content-hash filename to avoid duplicates.
    Returns the local file path, or None on failure.
    """
    response = requests.get(url, stream=True, timeout=15)
    response.raise_for_status()

    content = response.content
    if not content:
        logger.warning("Empty response body for image: %s", url)
        return None

    # Determine extension from Content-Type or URL
    content_type = response.headers.get("Content-Type", "")
    ext = mimetypes.guess_extension(content_type.split(";")[0].strip())
    if not ext or ext == ".bin":
        # Fallback: extract from URL
        parsed = urlparse(url)
        _, url_ext = os.path.splitext(parsed.path)
        ext = url_ext if url_ext else ".png"

    # Content-hash based filename to avoid duplicates
    content_hash = hashlib.md5(content).hexdigest()[:12]
    # Also preserve a readable portion of the original filename
    original_name = os.path.basename(urlparse(url).path).split("?")[0]
    # Remove non-alphanumeric chars except dots and underscores
    safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", original_name)
    if len(safe_name) > 40:
        safe_name = safe_name[:40]

    filename = f"{safe_name}_{content_hash}{ext}"
    dest_path = os.path.join(dest_dir, filename)

    # Skip if already exists (duplicate avoidance)
    if os.path.exists(dest_path):
        logger.debug("Image already exists locally, skipping: %s", dest_path)
        return dest_path

    with open(dest_path, "wb") as f:
        f.write(content)

    return dest_path


# ---------------------------------------------------------------------------
# B) DOCX image insertion with margin-respecting sizing
# ---------------------------------------------------------------------------

def add_image_to_doc(doc, image_source, max_width_inches=None):
    """
    Insert an image into a DOCX document, ensuring it fits within page margins.

    Args:
        doc: python-docx Document object.
        image_source: One of:
            - str: file path on disk
            - bytes: raw image data
            - BytesIO: image stream
        max_width_inches: Optional max width override. If None, uses the
            usable page width (page width minus margins).
    """
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_height = section.page_height - section.top_margin - section.bottom_margin

    if max_width_inches:
        max_width = Inches(max_width_inches)
    else:
        max_width = usable_width

    # Ensure max_width doesn't exceed usable width
    max_width = min(max_width, usable_width)

    try:
        tmp_path = None

        # Normalize image_source to a file path
        if isinstance(image_source, bytes):
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp_file.write(image_source)
            tmp_file.close()
            tmp_path = tmp_file.name
            file_path = tmp_path
        elif isinstance(image_source, io.BytesIO):
            image_source.seek(0)
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            tmp_file.write(image_source.read())
            tmp_file.close()
            tmp_path = tmp_file.name
            file_path = tmp_path
        else:
            file_path = str(image_source)

        if not os.path.exists(file_path):
            logger.warning("Image file not found: %s", file_path)
            doc.add_paragraph(f"[Image not found: {image_source}]")
            return

        # Get original dimensions and calculate scaled size
        with Image.open(file_path) as pil_img:
            img_width_px, img_height_px = pil_img.size

        # Convert pixels to EMU (1 inch = 96 px for screen, 914400 EMU)
        img_width_emu = int(img_width_px * 914400 / 96)
        img_height_emu = int(img_height_px * 914400 / 96)

        # Scale down if wider than max_width
        if img_width_emu > max_width:
            scale = max_width / img_width_emu
            img_width_emu = int(img_width_emu * scale)
            img_height_emu = int(img_height_emu * scale)

        # Also check height doesn't exceed usable page height
        if img_height_emu > usable_height:
            scale = usable_height / img_height_emu
            img_width_emu = int(img_width_emu * scale)
            img_height_emu = int(img_height_emu * scale)

        # Insert centered image
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(file_path, width=img_width_emu, height=img_height_emu)

    except Exception as e:
        logger.error("Error inserting image into DOCX: %s", e)
        doc.add_paragraph(f"[Error inserting image: {e}]")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def add_base64_image_to_doc_safe(doc, base64_str, ext):
    """
    Decode a base64 image string and insert it into the document
    respecting page margins. This is a drop-in replacement for the
    existing `add_base64_image_to_doc` function.
    """
    try:
        decoded_img = base64.b64decode(base64_str)
        add_image_to_doc(doc, decoded_img)
    except Exception as e:
        logger.error("Error decoding/inserting base64 image: %s", e)
        doc.add_paragraph(f"[Error processing base64 image: {e}]")


def resolve_image_path(image_src, writeup_name):
    """
    Resolve an image src (from HTML) to a local file path.

    Handles:
    - /protected_media/<writeup_name>/<filename>
    - /media/<filename>
    - Relative paths

    Returns:
        str or None: Absolute file path if found, else None.
    """
    if not image_src:
        return None

    # Handle protected_media paths
    if "/protected_media/" in image_src:
        filename = os.path.basename(image_src)
        safe_folder = writeup_name.replace(" ", "_")
        path = os.path.join(settings.PROTECTED_MEDIA_ROOT, safe_folder, filename)
        if os.path.exists(path):
            return path

    # Handle /media/ paths
    if image_src.startswith("/media/"):
        relative = image_src.replace("/media/", "")
        path = os.path.join(settings.MEDIA_ROOT, relative)
        if os.path.exists(path):
            return path

    # Handle absolute-looking paths starting with /
    if image_src.startswith("/"):
        path = os.path.join(settings.BASE_DIR, image_src.lstrip("/"))
        if os.path.exists(path):
            return path

    return None
